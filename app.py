import os
from flask import Flask, request, redirect, url_for, render_template, session, flash, send_file, jsonify, current_app, make_response
from dotenv import load_dotenv
import mysql.connector
from mysql.connector import Error
import json  # Adding this import to resolve the undefined variable errors
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.local import LocalProxy
from werkzeug.exceptions import HTTPException, InternalServerError
from urllib.parse import quote, urlparse
import logging
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
from datetime import datetime, timedelta
import secrets
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import smtplib
import requests
from itsdangerous import URLSafeTimedSerializer
from flask_wtf.csrf import CSRFProtect, generate_csrf
from flask_wtf import FlaskForm, RecaptchaField
from wtforms import StringField, PasswordField, SubmitField, BooleanField, IntegerField
from wtforms.validators import DataRequired, Email, EqualTo, InputRequired, Optional
from flask import send_from_directory
# from ai_recipe_generator import ai_bp # Import the blueprint
from blueprints.instructions.routes import bp as instructions_bp
from blueprints.quickstart.routes import bp as quickstart_bp # Import the quickstart blueprint
from blueprints.terms.routes import bp as terms_bp
from blueprints.privacy.routes import bp as privacy_bp
from blueprints.auth.routes import bp as auth_bp
from blueprints.auth.routes import LoginForm, ForgotPasswordForm, ResetPasswordForm, RegisterForm
from blueprints.auth.utils import generate_confirmation_token, confirm_token, send_confirmation_email, send_reset_email, hash_password
from blueprints import auth
from utils.db import get_db_connection
from blueprints.profile.routes import bp as profile_bp
from blueprints.about.routes import bp as about_bp  # Add this import
from forms import LeverancierForm, EenheidForm, CategorieForm, DishCategoryForm
from blueprints.haccp_api import bp as haccp_api_bp
import re  # Add this import
import html  # Add this import
import time 
from decimal import Decimal  # Add this import at the top
from routes.inventory import bp as inventory_bp
import boto3  # Add this for AWS S3 operations
from botocore.exceptions import ClientError  # Also add this for better error handling
import redis
from rq import Queue

load_dotenv()  # Load the values from .env

def create_app():
    """Application factory function"""
    app = Flask(__name__)
    
    # Move configuration from global scope to factory
    app.secret_key = os.getenv("SECRET_KEY")
    app.config['ENV'] = os.getenv('FLASK_ENV', 'production') # Set the environment
    app.config['SECURITY_PASSWORD_SALT'] = os.getenv("SECURITY_PASSWORD_SALT", "your-default-salt")
    app.config['RECAPTCHA_PUBLIC_KEY'] = os.getenv('RECAPTCHA_SITE_KEY')
    app.config['RECAPTCHA_PRIVATE_KEY'] = os.getenv('RECAPTCHA_SECRET_KEY')
    app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB 
    
    # RQ configuration (vervangt Celery configuratie)
    app.config['REDIS_URL'] = os.environ.get('REDIS_URL', 'redis://localhost:6379/0')
    
    app.config.update(
        SESSION_COOKIE_SECURE=True,
        SESSION_COOKIE_HTTPONLY=True,
        SESSION_COOKIE_SAMESITE='Lax',
        PERMANENT_SESSION_LIFETIME=1800
    )

    # Configure logging
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.DEBUG)
    ch = logging.StreamHandler()
    ch.setLevel(logging.DEBUG)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    ch.setFormatter(formatter)
    logger.addHandler(ch)

    # Initialize extensions
    csrf = CSRFProtect(app)
    app.wsgi_app = ProxyFix(app.wsgi_app)

    # Register blueprints
    app.register_blueprint(instructions_bp)
    app.register_blueprint(quickstart_bp, url_prefix='/quickstart')
    app.register_blueprint(terms_bp, url_prefix='/terms')
    app.register_blueprint(privacy_bp, url_prefix='/privacy')
    app.register_blueprint(auth_bp, template_folder='templates')
    app.register_blueprint(profile_bp)
    app.register_blueprint(about_bp)
    app.register_blueprint(haccp_api_bp)
    from blueprints.errors import bp as errors_bp
    app.register_blueprint(errors_bp)
    from blueprints.haccp.routes import bp as haccp_bp
    app.register_blueprint(haccp_bp, url_prefix='/haccp')
    app.register_blueprint(inventory_bp)
    
    # Import and register the ingredients_import blueprint
    from blueprints.ingredients_import import bp as ingredients_import_bp
    app.register_blueprint(ingredients_import_bp)
    
    # Import and register the suppliers blueprint
    from blueprints.suppliers import bp as suppliers_bp
    app.register_blueprint(suppliers_bp)
    
    # Import and register the takenboek blueprint
    from blueprints.takenboek import bp as takenboek_bp
    app.register_blueprint(takenboek_bp)
    
    # Import and register the ingredients blueprint
    from blueprints.ingredients import bp as ingredients_bp
    app.register_blueprint(ingredients_bp)

    # Import and register the admin blueprint
    from blueprints.admin import bp as admin_bp
    app.register_blueprint(admin_bp)
    
    # Register template filters and helper functions inside create_app
    def nl2br(value):
        return value.replace('\n', '<br>')
    app.template_filter('nl2br')(nl2br)

    # Custom decorator
    def login_required(f):
        def decorated_function(*args, **kwargs):
            if 'chef_id' not in session:
                flash("Geen toegang. Log opnieuw in.", "danger")
                return redirect(url_for('auth.login'))
            return f(*args, **kwargs)
        decorated_function.__name__ = f.__name__
        return decorated_function

    # Move all route handlers and helper functions inside create_app
    # Improved error handlers
    @app.errorhandler(Exception)
    def handle_exception(e):
        logger.error(f'Unhandled exception: {str(e)}')
        if isinstance(e, HTTPException):
            return render_template('errors/error.html', error=e), e.code
        
        error = InternalServerError()
        return render_template('errors/500.html', error=error), 500

    # Add middleware to redirect HTTP to HTTPS
    @app.before_request
    def before_request():
        if app.config['ENV'] == 'production':
            if request.url.startswith('http://'):
                url = request.url.replace('http://', 'https://', 1)
                code = 301
                return redirect(url, code=code)

    # Database configuration from .env
    DB_NAME = os.getenv("DB_NAME")
    DB_USER = os.getenv("DB_USER")
    DB_PASSWORD = os.getenv("DB_PASSWORD")
    DB_HOST = os.getenv("DB_HOST")
    DB_PORT = os.getenv("DB_PORT")

    # Database configuration - MODIFIED to prioritize local database
    db_url = os.getenv("JAWSDB_URL") 
    use_local_db = os.getenv("USE_LOCAL_DB", "true").lower() == "true"
    
    if db_url and not use_local_db:  # Only use JawsDB if explicitly not using local DB
        url = urlparse(db_url)
        DB_CONFIG = {
            'host': url.hostname,
            'database': url.path[1:],
            'user': url.username,
            'password': url.password,
            'port': url.port
        }
        print("Using production database (JawsDB)")
    else:  # Local development
        DB_CONFIG = {
            'host': os.getenv("DB_HOST"),
            'database': os.getenv("DB_NAME"),
            'user': os.getenv("DB_USER"),
            'password': os.getenv("DB_PASSWORD"),
            'port': os.getenv("DB_PORT")
        }
        print("Using local database")

    # Reset wachtwoord configuratie
    app.config['RESET_TOKEN_EXPIRE_MINUTES'] = 30
    app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER')
    app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
    app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
    app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
    app.config['MAIL_USE_TLS'] = True

    # Add the database connection function to the app context
    app.get_db_connection = get_db_connection

    # File storage configuration - ALWAYS use S3 regardless of environment
    app.config['USE_S3'] = True
    app.config['S3_BUCKET'] = os.getenv('S3_BUCKET')
    app.config['S3_KEY'] = os.getenv('AWS_ACCESS_KEY_ID')
    app.config['S3_SECRET'] = os.getenv('AWS_SECRET_ACCESS_KEY')
    app.config['S3_LOCATION'] = f"https://{app.config['S3_BUCKET']}.s3.amazonaws.com"
    
    from utils.storage import FileStorage
    storage = FileStorage(app)
    app.storage = storage
    
    # TEMPORARY DEBUGGING
    print(f"==== STORAGE CONFIGURATION ====")
    print(f"USE_S3: {app.config['USE_S3']}")
    print(f"S3_BUCKET: {app.config['S3_BUCKET']}")
    print(f"S3_LOCATION: {app.config['S3_LOCATION']}")
    print(f"Has S3_KEY: {'Yes' if app.config['S3_KEY'] else 'No'}")
    print(f"Has S3_SECRET: {'Yes' if app.config['S3_SECRET'] else 'No'}")
    print(f"==============================")
    
    # Use custom URL builder for static files
    @app.template_filter('file_url')
    def file_url_filter(path):
        if not path:
            return None
        return storage.get_file_url(path)

    # -----------------------------------------------------------
    #  Homepage (index)
    # -----------------------------------------------------------
    @app.route('/')
    def home():
        """
        Toon de homepage, met bijvoorbeeld een link naar inloggen/registreren.
        """
        return render_template('home.html', form=FlaskForm())  # Zorg voor een home.html template

    # -----------------------------------------------------------
    #  AVG Privacy Verklaring
    # -----------------------------------------------------------
    @app.route('/privacy')
    def privacy():
        return render_template('privacy.html', form=FlaskForm())

    # -----------------------------------------------------------
    #  Gerechten Samenstellen
    # -----------------------------------------------------------
    @app.route("/manage_dishes/<chef_naam>", methods=['GET', 'POST'])
    @login_required
    def manage_dishes(chef_naam):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        form = FlaskForm()
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('manage_dishes', chef_naam=chef_naam))
        
        cur = conn.cursor(dictionary=True)

        try:
            if request.method == 'POST' and request.form.get('gerechtForm') == '1':
                if not form.validate_on_submit():  # Add CSRF validation
                    flash("Ongeldige CSRF-token. Probeer het opnieuw.", "danger")
                    return redirect(url_for('manage_dishes', chef_naam=chef_naam))

                naam = request.form['naam']
                beschrijving = request.form['beschrijving']
                gerecht_categorie = request.form.get('gerecht_categorie')  # Use .get() to handle missing value
                ingredienten = request.form.get('ingredienten', '')  # Get ingredients from form and ensure it's not None
                bereidingswijze = request.form['bereidingswijze']

                cur.execute("""
                    INSERT INTO dishes (chef_id, naam, beschrijving, categorie, ingredienten, bereidingswijze)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (session['chef_id'], naam, beschrijving, gerecht_categorie, ingredienten, bereidingswijze))
                conn.commit()
                new_dish_id = cur.lastrowid

                flash('Gerecht succesvol aangemaakt!', 'success')
                return redirect(url_for('edit_dish', chef_naam=chef_naam, dish_id=new_dish_id))

            # Haal alle gerechten van deze chef op
            cur.execute("""
                SELECT * FROM dishes 
                WHERE chef_id = %s
                ORDER BY dish_id DESC
            """, (session['chef_id'],))
            alle_gerechten = cur.fetchall()

            # Haal alle gerecht categorieën op
            cur.execute("""
                SELECT * FROM dish_categories 
                WHERE chef_id = %s
                ORDER BY volgorde, naam
            """, (session['chef_id'],))
            dish_categories = cur.fetchall()

            return render_template('manage_dishes.html',
                                chef_naam=chef_naam,
                                gerechten=alle_gerechten,
                                dish_categories=dish_categories,
                                form=form)  # Add form to template context

        except Exception as e:
            logger.error(f'Error in manage_dishes: {str(e)}')
            flash("Er is een fout opgetreden.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
        
        finally:
            cur.close()
            conn.close()

    # -----------------------------------------------------------
    #  Gerecht Bewerken (Ingrediënten toevoegen)
    # -----------------------------------------------------------
    @app.route('/dashboard/<chef_naam>/dishes/<int:dish_id>', methods=['GET', 'POST'])
    @login_required
    def edit_dish(chef_naam, dish_id):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        form = FlaskForm()  # Add this line for CSRF validation
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        if request.method == 'POST':
            if not form.validate_on_submit():  # Add CSRF validation
                flash("Ongeldige CSRF-token. Probeer het opnieuw.", "danger")
                return redirect(url_for('edit_dish', chef_naam=chef_naam, dish_id=dish_id))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('manage_dishes', chef_naam=chef_naam))
        cur = conn.cursor(dictionary=True)

        # Haal info over het gerecht op (check of het bij deze chef hoort)
        cur.execute("SELECT * FROM dishes WHERE dish_id = %s AND chef_id = %s",
                    (dish_id, session['chef_id']))
        gerecht = cur.fetchone()
        if not gerecht:
            cur.close()
            conn.close()
            flash("Gerecht niet gevonden of je hebt geen toestemming.", "danger")
            return redirect(url_for('manage_dishes', chef_naam=chef_naam))

        # Haal dish_categories op
        cur.execute("""
            SELECT * FROM dish_categories 
            WHERE chef_id = %s
            ORDER BY volgorde, naam
        """, (session['chef_id'],))
        dish_categories = cur.fetchall()

        # Opslaan van nieuw ingrediënt voor dit gerecht
        if request.method == 'POST':
            if 'updateForm' in request.form:
                nieuwe_naam = request.form.get('naam')
                nieuwe_beschrijving = request.form.get('beschrijving')
                nieuwe_verkoopprijs = request.form.get('verkoopprijs')
                nieuwe_categorie = request.form.get('gerecht_categorie')
                nieuwe_ingredienten = request.form.get('ingredienten', '')  # Add default empty string
                nieuwe_bereidingswijze = request.form.get('bereidingswijze')

                try:
                    cur.execute("""
                        UPDATE dishes
                        SET naam = %s, beschrijving = %s, verkoopprijs = %s, categorie = %s, ingredienten = %s, bereidingswijze = %s
                        WHERE dish_id = %s AND chef_id = %s
                    """, (nieuwe_naam, nieuwe_beschrijving, nieuwe_verkoopprijs, nieuwe_categorie, nieuwe_ingredienten, nieuwe_bereidingswijze, dish_id, session['chef_id']))
                    conn.commit()
                    flash("Gerecht bijgewerkt!", "success")
                    return redirect(url_for('all_dishes'))
                except Exception as e:
                    conn.rollback()
                    flash(f"Fout bij bijwerken gerecht: {str(e)}", "danger")
            elif 'ingredientForm' in request.form:
                ingredient_id = request.form.get('ingredient_id')
                hoeveelheid = request.form.get('hoeveelheid')

                # Bepaal prijs_totaal (hoeveelheid * prijs_per_eenheid)
                cur.execute("""
                    SELECT prijs_per_eenheid 
                    FROM ingredients 
                    WHERE ingredient_id = %s 
                    AND chef_id = %s
                """, (ingredient_id, session['chef_id']))
                ingredient_info = cur.fetchone()
                if ingredient_info:
                    prijs_per_eenheid = ingredient_info['prijs_per_eenheid']
                    prijs_totaal = float(hoeveelheid) * float(prijs_per_eenheid)

                    try:
                        # Check if the ingredient already exists for this dish
                        cur.execute("""
                            SELECT * FROM dish_ingredients 
                            WHERE dish_id = %s AND ingredient_id = %s
                        """, (dish_id, ingredient_id))
                        existing_ingredient = cur.fetchone()

                        if existing_ingredient:
                            flash("Dit ingrediënt is al aan het gerecht toegevoegd.", "danger")
                        else:
                            cur.execute("""
                                INSERT INTO dish_ingredients (dish_id, ingredient_id, hoeveelheid, prijs_totaal)
                                VALUES (%s, %s, %s, %s)
                            """, (dish_id, ingredient_id, hoeveelheid, prijs_totaal))
                            conn.commit()
                            flash("Ingrediënt toegevoegd aan het gerecht!", "success")
                    except Exception as e:
                        conn.rollback()
                        flash(f"Fout bij toevoegen van ingrediënt: {str(e)}", "danger")
                else:
                    flash("Ongeldig ingrediënt of geen toegang.", "danger")
            elif 'descriptionForm' in request.form:
                nieuwe_beschrijving = request.form.get('beschrijving')
                try:
                    cur.execute("""
                        UPDATE dishes
                        SET beschrijving = %s
                        WHERE dish_id = %s AND chef_id = %s
                    """, (nieuwe_beschrijving, dish_id, session['chef_id']))
                    conn.commit()
                    flash("Beschrijving bijgewerkt!", "success")
                except Exception as e:
                    conn.rollback()
                    flash(f"Fout bij bijwerken beschrijving: {str(e)}", "danger")
            elif 'nameForm' in request.form:
                nieuwe_naam = request.form.get('naam')
                try:
                    cur.execute("""
                        UPDATE dishes
                        SET naam = %s
                        WHERE dish_id = %s AND chef_id = %s
                    """, (nieuwe_naam, dish_id, session['chef_id']))
                    conn.commit()
                    flash("Naam bijgewerkt!", "success")
                except Exception as e:
                    conn.rollback()
                    flash(f"Fout bij bijwerken naam: {str(e)}", "danger")
            elif 'priceForm' in request.form:
                nieuwe_verkoopprijs = request.form.get('verkoopprijs')
                try:
                    cur.execute("""
                        UPDATE dishes
                        SET verkoopprijs = %s
                        WHERE dish_id = %s AND chef_id = %s
                    """, (nieuwe_verkoopprijs, dish_id, session['chef_id']))
                    conn.commit()
                    flash("Verkoopprijs bijgewerkt!", "success")
                    return redirect(url_for('edit_dish', chef_naam=chef_naam, dish_id=dish_id) + '#verkoopprijs-sectie')
                except Exception as e:
                    conn.rollback()
                    flash(f"Fout bij bijwerken verkoopprijs: {str(e)}", "danger")
            elif 'methodForm' in request.form:
                nieuwe_bereidingswijze = request.form.get('bereidingswijze')
                try:
                    cur.execute("""
                        UPDATE dishes
                        SET bereidingswijze = %s
                        WHERE dish_id = %s AND chef_id = %s
                    """, (nieuwe_bereidingswijze, dish_id, session['chef_id']))
                    conn.commit()
                    flash("Bereidingswijze bijgewerkt!", "success")
                except Exception as e:
                    conn.rollback()
                    flash(f"Fout bij bijwerken bereidingswijze: {str(e)}", "danger")

        # Haal alle beschikbare ingrediënten van deze chef op
        cur.execute("""
            SELECT * FROM ingredients 
            WHERE chef_id = %s
            ORDER BY naam ASC
        """, (session['chef_id'],))
        alle_ingredienten = cur.fetchall()

        # Haal de gekoppelde ingrediënten voor dit gerecht op
        cur.execute("""
            SELECT di.*, i.naam AS ingredient_naam, i.eenheid, i.prijs_per_eenheid
            FROM dish_ingredients di
            JOIN ingredients i ON di.ingredient_id = i.ingredient_id
            WHERE di.dish_id = %s
        """, (dish_id,))
        gerecht_ingredienten = cur.fetchall()

        # Eventuele totale kostprijs (op basis van dish_ingredients.prijs_totaal)
        totaal_ingredient_prijs = sum([gi['prijs_totaal'] for gi in gerecht_ingredienten])

        # Haal alle beschikbare allergenen op
        cur.execute("SELECT * FROM allergenen ORDER BY naam")
        alle_allergenen = cur.fetchall()
        
        # Haal de geselecteerde allergenen voor dit gerecht op
        cur.execute("""
            SELECT allergeen_id FROM dish_allergenen
            WHERE dish_id = %s
        """, (dish_id,))
        gerecht_allergenen = [row['allergeen_id'] for row in cur.fetchall()]

        # Haal alle beschikbare diëten op
        cur.execute("SELECT * FROM dieten ORDER BY naam")
        alle_dieten = cur.fetchall()
        
        # Haal de geselecteerde diëten voor dit gerecht op
        cur.execute("""
            SELECT dieet_id FROM dish_dieten
            WHERE dish_id = %s
        """, (dish_id,))
        gerecht_dieten = [row['dieet_id'] for row in cur.fetchall()]

        cur.close()
        conn.close()

        return render_template(
            'edit_dish.html',
            chef_naam=chef_naam,
            gerecht=gerecht,
            alle_ingredienten=alle_ingredienten,
            gerecht_ingredienten=gerecht_ingredienten,
            totaal_ingredient_prijs=totaal_ingredient_prijs,
            alle_allergenen=alle_allergenen,
            gerecht_allergenen=gerecht_allergenen,
            alle_dieten=alle_dieten,
            gerecht_dieten=gerecht_dieten,  # Zorg dat deze correct wordt doorgegeven
            dish_categories=dish_categories, # Deze regel was al aanwezig, maar ik herhaal hem voor de zekerheid
            form=form  # Add form to template context
        )

    @app.route('/chef/<chef_naam>/dish/<int:dish_id>/ingredient/<int:ingredient_id>/update', methods=['POST'])
    @login_required
    def update_dish_ingredient(chef_naam, dish_id, ingredient_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        try:
            nieuwe_hoeveelheid = float(request.form['nieuwe_hoeveelheid'])

            conn = get_db_connection()
            if conn is None:
                return jsonify({'success': False, 'error': 'Database verbinding mislukt'}), 500

            cur = conn.cursor(dictionary=True)

            # Eerst de prijs_per_eenheid ophalen
            cur.execute("""
                SELECT prijs_per_eenheid 
                FROM ingredients 
                WHERE ingredient_id = %s 
                AND chef_id = %s
            """, (ingredient_id, session['chef_id']))
            ingredient = cur.fetchone()
            
            if ingredient:
                nieuwe_prijs_totaal = nieuwe_hoeveelheid * float(ingredient['prijs_per_eenheid'])

                # Update de hoeveelheid en prijs_totaal
                cur.execute("""
                    UPDATE dish_ingredients 
                    SET hoeveelheid = %s, prijs_totaal = %s
                    WHERE dish_id = %s AND ingredient_id = %s
                """, (nieuwe_hoeveelheid, nieuwe_prijs_totaal, dish_id, ingredient_id))
                
                # Update de totale ingrediëntprijs van het gerecht
                cur.execute("""
                    UPDATE dishes d
                    SET totaal_ingredient_prijs = (
                        SELECT COALESCE(SUM(di.prijs_totaal), 0)
                        FROM dish_ingredients di
                        WHERE di.dish_id = d.dish_id
                    )
                    WHERE d.dish_id = %s
                """, (dish_id,))
                
                conn.commit()
                return jsonify({'success': True})
            else:
                return jsonify({'success': False, 'error': 'Ingrediënt niet gevonden'}), 404

        except ValueError as e:
            return jsonify({'success': False, 'error': 'Ongeldige waarde opgegeven'}), 400
        except Exception as e:
            if conn:
                conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            if cur:
                cur.close()
            if conn:
                conn.close()

    @app.route('/chef/<chef_naam>/dish/<int:dish_id>/ingredient/<int:ingredient_id>/remove', methods=['POST'])
    @login_required
    def remove_dish_ingredient(chef_naam, dish_id, ingredient_id):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        form = FlaskForm()
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token.", "danger")
            return redirect(url_for('manage_dish_costs', chef_naam=chef_naam, dish_id=dish_id))  # Aangepaste redirect

        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('manage_dish_costs', chef_naam=chef_naam, dish_id=dish_id))  # Aangepaste redirect
        cur = conn.cursor()

        try:
            # Verwijder het ingredient uit het gerecht
            cur.execute("""
                DELETE FROM dish_ingredients 
                WHERE dish_id = %s AND ingredient_id = %s
            """, (dish_id, ingredient_id))
            
            conn.commit()
            flash("Ingrediënt verwijderd uit kostprijsberekening!", "success")
        except Exception as e:
            conn.rollback()
            flash(f"Fout bij verwijderen: {str(e)}", "danger")
        finally:
            cur.close()
            conn.close()

        return redirect(url_for('manage_dish_costs', chef_naam=chef_naam, dish_id=dish_id))  # Aangepaste redirect

    @app.route('/dashboard/<chef_naam>/dish/<int:dish_id>/allergenen', methods=['POST'])
    @login_required
    def update_dish_allergenen(chef_naam, dish_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        form = FlaskForm()  # Add CSRF validation
        if not form.validate_on_submit():
            return jsonify({'success': False, 'error': 'Ongeldige CSRF-token'}), 400

        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database connection error'}), 500
        cur = conn.cursor()

        try:
            # Verwijder bestaande allergenen voor dit gerecht
            cur.execute("DELETE FROM dish_allergenen WHERE dish_id = %s", (dish_id,))
            
            # Voeg nieuwe allergenen toe
            nieuwe_allergenen = request.form.getlist('allergenen[]')
            for allergeen_id in nieuwe_allergenen:
                cur.execute("""
                    INSERT INTO dish_allergenen (dish_id, allergeen_id)
                    VALUES (%s, %s)
                """, (dish_id, allergeen_id))
            
            conn.commit()
            return jsonify({'success': True, 'message': 'Allergenen bijgewerkt!'})
        except Exception as e:
            conn.rollback()
            app.logger.error(f"Fout bij bijwerken allergenen: {str(e)}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/dish/<int:dish_id>/dieten', methods=['POST'])
    @login_required
    def update_dish_dieten(chef_naam, dish_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        form = FlaskForm()
        if not form.validate_on_submit():
            return jsonify({'success': False, 'error': 'Ongeldige CSRF-token'}), 400

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database connection error'}), 500

        cur = conn.cursor(buffered=True)  # Gebruik buffered cursor
        max_retries = 3
        retry_count = 0

        while retry_count < max_retries:
            try:
                # Begin transaction
                cur.execute("SET TRANSACTION ISOLATION LEVEL READ COMMITTED")
                cur.fetchall()  # Lees resultaat van SET TRANSACTION
                
                # Lock the records in a consistent order to prevent deadlocks
                cur.execute("SELECT dish_id FROM dishes WHERE dish_id = %s FOR UPDATE", (dish_id,))
                cur.fetchall()  # Lees resultaat van SELECT FOR UPDATE
                
                # Delete existing diets
                cur.execute("DELETE FROM dish_dieten WHERE dish_id = %s", (dish_id,))
                
                # Add new diets
                dieten = request.form.getlist('dieten[]')
                if dieten:
                    # Use a single INSERT statement with multiple values
                    values = [(dish_id, int(dieet_id)) for dieet_id in dieten]
                    cur.executemany(
                        "INSERT INTO dish_dieten (dish_id, dieet_id) VALUES (%s, %s)",
                        values
                    )
                
                conn.commit()
                app.logger.debug("Diets updated successfully")
                return jsonify({'success': True})

            except Exception as e:
                conn.rollback()
                if "Deadlock" in str(e) and retry_count < max_retries - 1:
                    retry_count += 1
                    app.logger.warning(f"Deadlock encountered, retrying ({retry_count}/{max_retries})")
                    time.sleep(0.1 * retry_count)  # Exponential backoff
                    continue
                
                app.logger.error(f"Error updating diets: {str(e)}")
                return jsonify({'success': False, 'error': str(e)}), 500
            finally:
                cur.close()
                conn.close()

    # -----------------------------------------------------------
    #  Alle Gerechten Beheren
    # -----------------------------------------------------------
    @app.route('/all_dishes')
    @login_required
    def all_dishes():
        """
        Pagina om alle gerechten te beheren.
        """
        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('home'))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal eerst de dish categories op
            cur.execute("""
                SELECT * FROM dish_categories 
                WHERE chef_id = %s
                ORDER BY volgorde, naam
            """, (session['chef_id'],))
            dish_categories = cur.fetchall()

            # Aangepaste query om alleen gerechten van de ingelogde chef te tonen
            cur.execute("""
                SELECT d.*, c.naam as chef_naam, 
                    (SELECT SUM(di.prijs_totaal) 
                        FROM dish_ingredients di 
                        WHERE di.dish_id = d.dish_id) as totaal_ingredient_prijs
                FROM dishes d
                JOIN chefs c ON d.chef_id = c.chef_id
                WHERE d.chef_id = %s  
                ORDER BY d.dish_id DESC
            """, (session['chef_id'],))
            alle_gerechten = cur.fetchall()

            return render_template('all_dishes.html', 
                                gerechten=alle_gerechten,
                                dish_categories=dish_categories,
                                form=FlaskForm())

        except Exception as e:
            flash(f"Error: {str(e)}", "danger")
            return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
        finally:
            cur.close()
            conn.close()

    # -----------------------------------------------------------
    #  Nieuw Gerecht (alleen naam)
    # -----------------------------------------------------------
    @app.route('/create_dish', methods=['GET', 'POST'])
    @login_required
    def create_dish():
        """
        Route om een nieuw gerecht aan te maken met alleen de naam.
        """
        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        form = FlaskForm()
        if request.method == 'POST':
            if not form.validate_on_submit():
                flash("Ongeldige CSRF-token.", "danger")
                return redirect(url_for('create_dish'))

            naam = request.form.get('naam')
            if not naam:
                flash("Naam is verplicht.", "danger")
                return redirect(url_for('create_dish'))

            conn = get_db_connection()
            if conn is None:
                flash("Database connection error.", "danger")
                return redirect(url_for('all_dishes'))
            cur = conn.cursor()

            try:
                cur.execute("""
                    INSERT INTO dishes (chef_id, naam)
                    VALUES (%s, %s)
                """, (session['chef_id'], naam))
                conn.commit()
                new_dish_id = cur.lastrowid
                flash("Gerecht succesvol aangemaakt! U kunt nu de overige gegevens invullen.", "success")
                return redirect(url_for('edit_dish', chef_naam=session['chef_naam'], dish_id=new_dish_id))
            except Exception as e:
                conn.rollback()
                flash(f"Fout bij aanmaken gerecht: {str(e)}", "danger")
                return redirect(url_for('all_dishes'))
            finally:
                cur.close()
                conn.close()

        return render_template('create_dish.html', form=form)

    # -----------------------------------------------------------
    #  Export Dishes to MS Word
    # -----------------------------------------------------------
    @app.route('/export_dishes', methods=['POST'])
    @login_required
    def export_dishes():
        """
        Export selected dishes to a Microsoft Word document.
        """
        form = FlaskForm()
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token.", "danger")
            return redirect(url_for('all_dishes'))

        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        selected_dish_ids = request.form.getlist('selected_dishes')
        if not selected_dish_ids:
            flash("Geen gerechten geselecteerd.", "danger")
            return redirect(url_for('all_dishes'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('all_dishes'))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal alleen de geselecteerde gerechten van de ingelogde chef op
            format_strings = ','.join(['%s'] * len(selected_dish_ids))
            query = """
                SELECT d.*, c.naam as chef_naam, 
                    (SELECT SUM(di.prijs_totaal) 
                        FROM dish_ingredients di 
                        WHERE di.dish_id = d.dish_id) as totaal_ingredient_prijs
                FROM dishes d
                JOIN chefs c ON d.chef_id = c.chef_id
                WHERE d.dish_id IN ({}) AND d.chef_id = %s
                ORDER BY d.categorie
            """.format(format_strings)
            params = selected_dish_ids + [session['chef_id']]
            cur.execute(query, tuple(params))
            selected_dishes = cur.fetchall()

            # Maak een Word-document aan
            doc = Document()
            doc.add_heading('Menukaart', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Groepeer gerechten per categorie
            current_category = None
            for dish in selected_dishes:
                # Voeg categorieheader toe als we een nieuwe categorie tegenkomen
                if dish['categorie'] != current_category:
                    current_category = dish['categorie']
                    doc.add_heading(current_category, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Voeg gerechtnaam en prijs toe
                verkoopprijs = dish['verkoopprijs'] if dish['verkoopprijs'] else 'n.v.t.'
                naam = dish['naam'] if dish['naam'] else 'Onbekend gerecht'
                price_paragraph = doc.add_paragraph()
                price_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                price_paragraph.add_run(f"{naam} - €{verkoopprijs}").bold = True

                # Voeg beschrijving toe en converteer van HTML naar tekst
                if dish['beschrijving']:
                    desc_text = html_to_text(dish['beschrijving'])
                    desc_paragraph = doc.add_paragraph(desc_text)
                    desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Haal allergenen op voor dit gerecht
                cur.execute("""
                    SELECT a.naam, a.icon_class 
                    FROM allergenen a
                    JOIN dish_allergenen da ON a.allergeen_id = da.allergeen_id
                    WHERE da.dish_id = %s
                    ORDER BY a.naam
                """, (dish['dish_id'],))
                allergenen = cur.fetchall()

                # Voeg allergenen toe als ze bestaan
                if allergenen:
                    allergenen_paragraph = doc.add_paragraph()
                    allergenen_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    allergenen_text = "Allergenen: " + ", ".join([f"{a['icon_class']} {a['naam']}" for a in allergenen])
                    allergenen_paragraph.add_run(allergenen_text).italic = True

                # Voeg diëten toe aan het document
                cur.execute("""
                    SELECT d.naam, d.icon_class 
                    FROM dieten d
                    JOIN dish_dieten dd ON d.dieet_id = dd.dieet_id
                    WHERE dd.dish_id = %s
                    ORDER BY d.naam
                """, (dish['dish_id'],))
                dieten = cur.fetchall()
                
                if dieten:
                    dieten_paragraph = doc.add_paragraph()
                    dieten_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    dieten_text = "Geschikt voor: " + ", ".join([f"{d['icon_class']} {d['naam']}" for d in dieten])
                    dieten_paragraph.add_run(dieten_text).italic = True

                # Voeg witruimte toe tussen gerechten
                doc.add_paragraph()

            # Sla het document op in een in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(
                buffer,
                as_attachment=True,
                download_name='Menukaart.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            logger.error(f'Error exporting dishes: {str(e)}')
            flash("Er is een fout opgetreden bij het exporteren van de menukaart.", "danger")
            return redirect(url_for('all_dishes'))
        finally:
            cur.close()
            conn.close()

    # -----------------------------------------------------------
    #  Export Cookbook to MS Word
    # -----------------------------------------------------------
    @app.route('/export_cookbook', methods=['POST'])
    @login_required
    def export_cookbook():
        form = FlaskForm()  # Add CSRF validation
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token. Probeer het opnieuw.", "danger")
            return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))

        if 'chef_id' not in session:  # Add this check
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('all_dishes'))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal alle gerechten van de ingelogde chef op
            cur.execute("""
                SELECT d.*, c.naam as chef_naam, 
                    (SELECT SUM(di.prijs_totaal) 
                        FROM dish_ingredients di 
                        WHERE di.dish_id = d.dish_id) as totaal_ingredient_prijs
                FROM dishes d
                JOIN chefs c ON d.chef_id = c.chef_id
                WHERE d.chef_id = %s
                ORDER BY CASE 
                    WHEN d.categorie = 'Amuse-bouche' THEN 1
                    WHEN d.categorie = 'Hors-d''oeuvre' THEN 2
                    WHEN d.categorie = 'Potage' THEN 3
                    WHEN d.categorie = 'Poisson' THEN 4
                    WHEN d.categorie = 'Entrée' THEN 5
                    WHEN d.categorie = 'Sorbet' THEN 6
                    WHEN d.categorie = 'Relevé of Rôti' THEN 7
                    WHEN d.categorie = 'Légumes / Groentegerecht' THEN 8
                    WHEN d.categorie = 'Salade' THEN 9
                    WHEN d.categorie = 'Fromage' THEN 10
                    WHEN d.categorie = 'Entremets' THEN 11
                    WHEN d.categorie = 'Café / Mignardises' THEN 12
                    WHEN d.categorie = 'Digestief' THEN 13
                    ELSE 14
                END
            """, (session['chef_id'],))
            all_dishes = cur.fetchall()

            # Maak een Word-document aan
            doc = Document()
            doc.add_heading('Kookboek', 0).alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Voeg een inhoudsopgave toe
            doc.add_heading('Inhoudsopgave', level=1)
            for dish in all_dishes:
                doc.add_paragraph(dish['naam'], style='List Number')

            doc.add_page_break()

            # Voeg veiligheidscontroles toe voor alle velden
            for dish in all_dishes:
                verkoopprijs = dish['verkoopprijs'] if dish['verkoopprijs'] else 'n.v.t.'
                naam = dish['naam'] if dish['naam'] else 'Onbekend gerecht'
                heading = doc.add_heading(f"{naam} - €{verkoopprijs}", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                description_text = html_to_text(dish['beschrijving']) if dish['beschrijving'] else 'Geen beschrijving'
                description = doc.add_paragraph(description_text)
                description.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Haal de ingrediënten voor dit gerecht op
                cur.execute("""
                    SELECT di.hoeveelheid, i.naam AS ingredient_naam, i.eenheid
                    FROM dish_ingredients di
                    JOIN ingredients i ON di.ingredient_id = i.ingredient_id
                    WHERE di.dish_id = %s
                    ORDER BY i.naam ASC
                """, (dish['dish_id'],))
                gerecht_ingredienten = cur.fetchall()

                if gerecht_ingredienten:
                    doc.add_heading('Ingrediënten', level=2)
                    for gi in gerecht_ingredienten:
                        ingredient_text = f"{gi['hoeveelheid']} {gi['eenheid']} {gi['ingredient_naam']}"
                        doc.add_paragraph(ingredient_text).alignment = WD_ALIGN_PARAGRAPH.LEFT

                method_text = html_to_text(dish['bereidingswijze']) if dish['bereidingswijze'] else 'Geen bereidingswijze'
                doc.add_heading('Bereidingswijze', level=2)
                method = doc.add_paragraph(method_text)
                method.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Haal allergenen op voor dit gerecht
                cur.execute("""
                    SELECT a.naam, a.icon_class 
                    FROM allergenen a
                    JOIN dish_allergenen da ON a.allergeen_id = da.allergeen_id
                    WHERE da.dish_id = %s
                    ORDER BY a.naam
                """, (dish['dish_id'],))
                allergenen = cur.fetchall()

                # Voeg allergenen toe als ze bestaan
                if allergenen:
                    allergenen_paragraph = doc.add_paragraph()
                    allergenen_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    allergenen_text = "Allergenen: " + ", ".join([f"{a['icon_class']} {a['naam']}" for a in allergenen])
                    allergenen_paragraph.add_run(allergenen_text).italic = True

                # Voeg diëten toe aan het document
                cur.execute("""
                    SELECT d.naam, d.icon_class 
                    FROM dieten d
                    JOIN dish_dieten dd ON d.dieet_id = dd.dieet_id
                    WHERE dd.dish_id = %s
                    ORDER BY d.naam
                """, (dish['dish_id'],))
                dieten = cur.fetchall()
                
                if dieten:
                    dieten_paragraph = doc.add_paragraph()
                    dieten_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    dieten_text = "Geschikt voor: " + ", ".join([f"{d['icon_class']} {d['naam']}" for d in dieten])
                    dieten_paragraph.add_run(dieten_text).italic = True

                doc.add_paragraph("\n").alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Sla het document op in een in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            safe_filename = quote('Kookboek.docx', safe='')
            return send_file(
                buffer,
                as_attachment=True,
                download_name=safe_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            logger.error(f'Error exporting cookbook: {str(e)}')
            flash("Er is een fout opgetreden bij het exporteren van het kookboek.", "danger")
            return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
        finally:
            cur.close()
            conn.close()

    # -----------------------------------------------------------
    #  Verwijder Gerecht
    # -----------------------------------------------------------
    @app.route('/delete_dish/<int:dish_id>', methods=['POST'])
    @login_required
    def delete_dish(dish_id):
        form = FlaskForm()
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token.", "danger")
            return redirect(url_for('all_dishes'))

        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        """
        Verwijder een gerecht.
        """
        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('all_dishes'))
        cur = conn.cursor()

        try:
            cur.execute("DELETE FROM dish_ingredients WHERE dish_id = %s", (dish_id,))
            cur.execute("DELETE FROM dishes WHERE dish_id = %s AND chef_id = %s", (dish_id, session['chef_id']))
            conn.commit()
            flash("Gerecht succesvol verwijderd!", "success")
        except Exception as e:
            conn.rollback()
            flash(f"Fout bij verwijderen gerecht: {str(e)}", "danger")
        finally:
            cur.close()
            conn.close()

        return redirect(url_for('all_dishes'))

    # -----------------------------------------------------------
    #  Export Dish to MS Word
    # -----------------------------------------------------------
    def html_to_text(html_content):
        """Convert HTML content to plain text by removing HTML tags"""
        if not html_content:
            return ''
        try:
            html_content = re.sub(r'<style[^>]*>[\s\S]*?</style>', '', html_content)
        except (TypeError, AttributeError):
            html_content = ''
        # Replace common line break elements with newlines first
        html_content = re.sub(r'<br\s*/?>', '\n', html_content)
        html_content = re.sub(r'</p><p>', '\n\n', html_content)
        html_content = re.sub(r'</li><li>', '\n', html_content)
        html_content = re.sub(r'</ul>', '\n', html_content)
        # Remove all HTML tags
        text = re.sub(r'<[^>]+>', '', html_content)
        # Convert HTML entities
        text = html.unescape(text)
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        # Preserve line breaks (don't collapse them)
        text = re.sub(r'\n+', '\n', text)
        return text.strip()

    def format_ingredients_list(ingredients_text):
        """Special handling for ingredients to ensure each item appears on a new line"""
        if not ingredients_text:
            return ''
        
        # First convert from HTML to text
        text = html_to_text(ingredients_text)
        
        # Try to detect if the text is already formatted with line breaks
        if '\n' in text:
            return text  # Already has line breaks, return as is
            
        # Split by common ingredient list separators if no line breaks found
        items = []
        # Try splitting by semicolons first
        if ';' in text:
            items = [item.strip() for item in text.split(';')]
        # Otherwise check for numbered items or bullet points
        elif bool(re.search(r'^\d+\.|\d+\)', text)):
            # Split by numbered items (1. or 1) pattern)
            items = re.split(r'(?:\d+\.|\d+\))\s*', text)
            items = [item.strip() for item in items if item.strip()]
        # Or split by comma if items likely contain amounts (numbers with units)
        elif ',' in text and bool(re.search(r'\d+\s*(g|kg|l|ml|cl|el|tl|st)', text)):
            items = [item.strip() for item in text.split(',')]
        
        # If we identified list items, join them with new lines
        if items:
            return '\n'.join(items)
        
        # Otherwise return original text with preserved line breaks
        return text

    @app.template_filter('format_price')
    def format_price(value):
        """Custom filter to format prices with 5 decimal places"""
        try:
            return '{:.5f}'.format(float(value))
        except (ValueError, TypeError):
            return '0.00000'

    @app.route('/export_dish/<chef_naam>/<dish_id>', methods=['POST'])
    @login_required
    def export_dish(chef_naam, dish_id):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        form = FlaskForm()
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token.", "danger")
            return redirect(url_for('all_dishes'))

        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('manage_dishes', chef_naam=chef_naam))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal gerecht informatie op
            cur.execute("""
                SELECT * FROM dishes 
                WHERE dish_id = %s AND chef_id = %s
            """, (dish_id, session['chef_id']))
            gerecht = cur.fetchone()
            
            if not gerecht:
                flash("Gerecht niet gevonden of geen toegang.", "danger")
                return redirect(url_for('all_dishes'))

            # Haal ingrediënten op voor kostprijsberekening
            cur.execute("""
                SELECT di.hoeveelheid, di.prijs_totaal, i.naam AS ingredient_naam, i.eenheid, i.prijs_per_eenheid
                FROM dish_ingredients di
                JOIN ingredients i ON di.ingredient_id = i.ingredient_id
                WHERE di.dish_id = %s
            """, (dish_id,))
            gerecht_ingredienten = cur.fetchall()
            
            # Bereken totale kostprijs
            totaal_ingredient_prijs = sum(float(ingredient['prijs_totaal']) for ingredient in gerecht_ingredienten)

            # Maak Word document
            doc = Document()
            doc.add_heading(f"{gerecht['naam']} - Receptuur", level=1)
            
            if gerecht['beschrijving']:
                doc.add_heading('Beschrijving', level=2)
                doc.add_paragraph(html_to_text(gerecht['beschrijving']))
            
            # Gebruik ingredienten uit het tekstveld in plaats van de tabel
            doc.add_heading('Ingrediënten', level=2)
            if gerecht['ingredienten'] and gerecht['ingredienten'].strip():
                # Apply special formatting for ingredients list to preserve structure
                ingredients_text = format_ingredients_list(gerecht['ingredienten'])
                doc.add_paragraph(ingredients_text)
            else:
                doc.add_paragraph("Geen ingrediëntenlijst opgegeven.")

            if gerecht['bereidingswijze']:
                doc.add_heading('Bereidingswijze', level=2)
                doc.add_paragraph(html_to_text(gerecht['bereidingswijze']))
                
            # Voeg kostprijsberekening toe onderaan het document
            doc.add_heading('Kostprijsberekening', level=2)
            price_table = doc.add_table(rows=1, cols=4)
            price_table.style = 'Table Grid'
            
            # Headers voor kostprijstabel
            price_header = price_table.rows[0].cells
            price_header[0].text = 'Ingrediënt'
            price_header[1].text = 'Hoeveelheid'
            price_header[2].text = 'Prijs per eenheid'
            price_header[3].text = 'Totaal'
            
            # Maak de headers vet
            for cell in price_header:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Voeg ingrediënten met prijzen toe
            for ingredient in gerecht_ingredienten:
                row_cells = price_table.add_row().cells
                row_cells[0].text = ingredient['ingredient_naam']
                row_cells[1].text = f"{ingredient['hoeveelheid']} {ingredient['eenheid']}"
                row_cells[2].text = f"€{ingredient['prijs_per_eenheid']}"
                row_cells[3].text = f"€{ingredient['prijs_totaal']}"
            
            # Voeg totaal toe onderaan de tabel
            total_row = price_table.add_row().cells
            total_row[0].text = ''
            total_row[1].text = ''
            total_row[2].text = 'Totale kostprijs:'
            total_row[3].text = f"€{totaal_ingredient_prijs:.2f}"
            
            # Maak de totaalrij vet
            for i in range(2, 4):
                for paragraph in total_row[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Voeg verkoopprijs toe onder de kostprijstabel als deze beschikbaar is
            if gerecht['verkoopprijs']:
                price_paragraph = doc.add_paragraph()
                price_paragraph.add_run(f"Verkoopprijs: €{gerecht['verkoopprijs']}").bold = True
                
                # Bereken en toon winstmarge als verkoopprijs aanwezig is
                if float(gerecht['verkoopprijs']) > 0:
                    margin = float(gerecht['verkoopprijs']) - totaal_ingredient_prijs
                    margin_percentage = (margin / float(gerecht['verkoopprijs'])) * 100 if float(gerecht['verkoopprijs']) > 0 else 0
                    margin_paragraph = doc.add_paragraph()
                    margin_paragraph.add_run(f"Winstmarge: €{margin:.2f} ({margin_percentage:.2f}%)").italic = True

            # Sla document op in memory
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            filename = f"{gerecht['naam']}_recept.docx"
            safe_filename = secure_filename(filename)
            return send_file(
                buffer,
                as_attachment=True,
                download_name=safe_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            logger.error(f'Error exporting dish: {str(e)}')
            flash("Er is een fout opgetreden bij het exporteren van het recept.", "danger")
            return redirect(url_for('all_dishes'))
        finally:
            cur.close()
            conn.close()

    # -----------------------------------------------------------
    #  Backup & Restore Page
    # -----------------------------------------------------------
    @app.route('/dashboard/<chef_naam>/backup-restore')
    @login_required
    def backup_restore(chef_naam):
        """Display backup and restore page"""
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))
            
        return render_template('backup_restore.html', chef_naam=chef_naam)

    # -----------------------------------------------------------
    #  Werkinstructie
    # -----------------------------------------------------------
    # @app.route('/instructions')
    # def instructions():
    #     return render_template('instructions.html', form=FlaskForm())

    # -----------------------------------------------------------
    #  Bestellijst Beheren
    # -----------------------------------------------------------
    @app.route('/orderlist', methods=['GET', 'POST'])
    @login_required
    def manage_orderlist():
        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('home'))
        cur = conn.cursor(dictionary=True)

        # Aangepaste query om alleen gerechten van de ingelogde chef te tonen
        cur.execute("""
            SELECT d.*, c.naam as chef_naam
            FROM dishes d
            JOIN chefs c ON d.chef_id = c.chef_id
            WHERE d.chef_id = %s
            ORDER BY d.naam
        """, (session['chef_id'],))
        alle_gerechten = cur.fetchall()
        cur.close()
        conn.close()

        return render_template('manage_orderlist.html', gerechten=alle_gerechten, form=FlaskForm())

    @app.route('/export_orderlist', methods=['POST'])
    @login_required
    def export_orderlist():
        form = FlaskForm()  # Add CSRF validation
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token. Probeer het opnieuw.", "danger")
            return redirect(url_for('manage_orderlist'))

        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        dish_quantities = {}
        for key, value in request.form.items():
            if key.startswith('quantity_') and value and int(value) > 0:
                dish_id = int(key.replace('quantity_', ''))
                dish_quantities[dish_id] = int(value)

        if not dish_quantities:
            flash("Geen gerechten geselecteerd.", "warning")
            return redirect(url_for('manage_orderlist'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('manage_orderlist'))
        cur = conn.cursor(dictionary=True)

        try:
            total_cost = 0  # Variabele voor totale kostprijs
            ingredient_total = 0  # Initialize here

            # Gebruik een dictionary met tuple als key voor de ingrediënten
            total_ingredients = {}
            dishes_info = {}

            for dish_id, quantity in dish_quantities.items():
                # Haal gerecht informatie op
                cur.execute("""
                    SELECT * FROM dishes 
                    WHERE dish_id = %s
                """, (dish_id,))
                dish = cur.fetchone()
                if dish:  # Voeg veiligheidscheck toe
                    dishes_info[dish_id] = {
                        'naam': dish['naam'],
                        'aantal': quantity,
                        'verkoopprijs': float(dish['verkoopprijs'] or 0)
                    }
                    # Bereken totale verkoopprijs
                    total_cost += dishes_info[dish_id]['verkoopprijs'] * quantity

                    # Haal ingrediënten voor dit gerecht op met foutafhandeling
                    try:
                        cur.execute("""
                            SELECT di.hoeveelheid, 
                                i.naam AS ingredient_naam, 
                                i.eenheid, 
                                i.prijs_per_eenheid,
                                l.leverancier_id,
                                l.naam AS leverancier_naam
                            FROM dish_ingredients di
                            JOIN ingredients i ON di.ingredient_id = i.ingredient_id
                            LEFT JOIN leveranciers l ON i.leverancier_id = l.leverancier_id
                            WHERE di.dish_id = %s
                        """, (dish_id,))
                        ingredients = cur.fetchall()
                        for ing in ingredients:
                            # Calculate ingredient subtotal and add to total
                            ingredient_subtotal = float(ing['hoeveelheid']) * quantity * float(ing['prijs_per_eenheid'] or 0)
                            ingredient_total += ingredient_subtotal  # Add to running total
                            
                            supplier_key = ing['leverancier_naam'] or 'Geen leverancier'
                            if supplier_key not in total_ingredients:
                                total_ingredients[supplier_key] = []
                            total_ingredients[supplier_key].append({
                                'naam': ing['ingredient_naam'],
                                'hoeveelheid': float(ing['hoeveelheid']) * quantity,
                                'eenheid': ing['eenheid'],
                                'prijs': float(ing['prijs_per_eenheid'] or 0)
                            })
                    except Exception as e:
                        logger.error(f'Error processing ingredients for dish {dish_id}: {str(e)}')
                        continue

            # Maak Word document
            doc = Document()
            doc.add_heading('Bestellijst per Leverancier', 0)

            for supplier, ingredients in total_ingredients.items():
                doc.add_heading(f'Leverancier: {supplier}', level=1)
                table = doc.add_table(rows=1, cols=5)  # Verhoogd naar 5 kolommen
                table.style = 'Table Grid'
                
                # Voeg headers toe aan eerste rij
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Ingredient'
                header_cells[1].text = 'Aantal'
                header_cells[2].text = 'Eenheid'
                header_cells[3].text = 'Prijs per eenheid'
                header_cells[4].text = 'Totaal'  # Nieuwe kolom
                
                # Maak headers bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Voeg ingrediënten toe met totaalprijs
                for ingredient in ingredients:
                    row = table.add_row().cells
                    row[0].text = ingredient['naam']
                    row[1].text = f"{ingredient['hoeveelheid']:.0f}"  # Gehele getallen voor hoeveelheden
                    row[2].text = ingredient['eenheid']
                    row[3].text = f"€{ingredient['prijs']:.5f}"       # 5 decimalen voor prijzen
                    # Bereken en voeg totaalprijs toe
                    totaal = ingredient['hoeveelheid'] * ingredient['prijs']
                    row[4].text = f"€{totaal:.5f}"                    # 5 decimalen voor totaalprijzen

                doc.add_paragraph()

            # Voeg financieel overzicht toe
            doc.add_heading('Financieel Overzicht:', level=1)
            
            # Ingrediëntenkosten
            cost_table = doc.add_table(rows=1, cols=2)
            cost_table.style = 'Table Grid'
            cost_row = cost_table.rows[0].cells
            cost_row[0].text = 'Totale Ingrediëntenkosten:'
            cost_row[1].text = f"€{ingredient_total:.5f}"

            # Verwachte verkoopprijs
            price_row = cost_table.add_row().cells
            price_row[0].text = 'Verwachte Verkoopprijs:'
            price_row[1].text = f"€{total_cost:.5f}"

            # Verwachte winst
            expected_profit = total_cost - ingredient_total
            profit_row = cost_table.add_row().cells
            profit_row[0].text = 'Verwachte Winst:'
            profit_row[1].text = f"€{expected_profit:.5f}"

            # Exporteer document
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(
                buffer,
                as_attachment=True,
                download_name='bestellijst.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            logger.error(f'Error creating orderlist: {str(e)}')
            flash("Er is een fout opgetreden bij het maken van de bestellijst.", "danger")
            return redirect(url_for('manage_orderlist'))
        finally:
            cur.close()
            conn.close()

    # -----------------------------------------------------------
    #  HACCP Module
    # -----------------------------------------------------------
    @app.route('/dashboard/<chef_naam>/haccp')
    @login_required
    def haccp_dashboard(chef_naam):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
        cur = conn.cursor(dictionary=True)

        # Haal HACCP checklists op met extra informatie
        cur.execute("""
            SELECT c.*, 
                (SELECT COUNT(*) 
                    FROM haccp_metingen m 
                    JOIN haccp_checkpunten p ON m.punt_id = p.punt_id 
                    WHERE p.checklist_id = c.checklist_id 
                    AND DATE(m.timestamp) = CURDATE()) as metingen_vandaag,
                DATEDIFF(CURDATE(), 
                        IFNULL((SELECT MAX(DATE(m.timestamp))
                                FROM haccp_metingen m
                                JOIN haccp_checkpunten p ON m.punt_id = p.punt_id
                                WHERE p.checklist_id = c.checklist_id), 
                                '1970-01-01')) as dagen_sinds_laatste_meting
            FROM haccp_checklists c
            WHERE c.chef_id = %s
            ORDER BY c.created_at DESC
        """, (session['chef_id'],))
        checklists = cur.fetchall()

        # Haal laatste metingen op
        cur.execute("""
            SELECT m.*, c.omschrijving, c.grenswaarde
            FROM haccp_metingen m
            JOIN haccp_checkpunten c ON m.punt_id = c.punt_id
            WHERE m.chef_id = %s
            ORDER BY m.timestamp DESC
            LIMIT 10
        """, (session['chef_id'],))
        laatste_metingen = cur.fetchall()

        cur.close()
        conn.close()

        return render_template('haccp/dashboard.html',
                            chef_naam=chef_naam,
                            checklists=checklists,
                            laatste_metingen=laatste_metingen)

    @app.route('/dashboard/<chef_naam>/haccp/new_checklist', methods=['GET', 'POST'])
    @login_required
    def new_haccp_checklist(chef_naam):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        if request.method == 'POST':
            naam = request.form.get('naam')
            frequentie = request.form.get('frequentie')
            
            conn = get_db_connection()
            if conn is None:
                flash("Database connection error.", "danger")
                return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))
            cur = conn.cursor()

            try:
                cur.execute("""
                    INSERT INTO haccp_checklists (chef_id, naam, frequentie)
                    VALUES (%s, %s, %s)
                """, (session['chef_id'], naam, frequentie))
                checklist_id = cur.lastrowid
                
                # Voeg checkpunten toe
                checkpunten = request.form.getlist('checkpunt[]')
                grenswaarden = request.form.getlist('grenswaarde[]')
                acties = request.form.getlist('actie[]')
                
                for i in range(len(checkpunten)):
                    cur.execute("""
                        INSERT INTO haccp_checkpunten 
                        (checklist_id, omschrijving, grenswaarde, corrigerende_actie)
                        VALUES (%s, %s, %s, %s)
                    """, (checklist_id, checkpunten[i], grenswaarden[i], acties[i]))
                
                conn.commit()
                flash("HACCP-checklist aangemaakt!", "success")
                
            except Exception as e:
                conn.rollback()
                flash(f"Fout bij aanmaken checklist: {str(e)}", "danger")
            finally:
                cur.close()
                conn.close()
                return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))

        return render_template('haccp/new_checklist.html', chef_naam=chef_naam, form=FlaskForm())

    @app.route('/dashboard/<chef_naam>/haccp/checklist/<int:checklist_id>/fill', methods=['GET', 'POST'])
    @login_required
    def fill_haccp_checklist(chef_naam, checklist_id):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal checklist informatie op
            cur.execute("""
                SELECT * FROM haccp_checklists 
                WHERE checklist_id = %s AND chef_id = %s
            """, (checklist_id, session['chef_id']))
            checklist = cur.fetchone()

            if not checklist:
                flash("Checklist niet gevonden.", "danger")
                return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))

            # Haal alle checkpunten op
            cur.execute("""
                SELECT * FROM haccp_checkpunten 
                WHERE checklist_id = %s
            """, (checklist_id,))
            checkpunten = cur.fetchall()

            if request.method == 'POST':
                for punt in checkpunten:
                    waarde = request.form.get(f'waarde_{punt["punt_id"]}')
                    opmerking = request.form.get(f'opmerking_{punt["punt_id"]}')
                    actie = request.form.get(f'actie_{punt["punt_id"]}')

                    if waarde:  # Alleen opslaan als er een waarde is ingevuld
                        cur.execute("""
                            INSERT INTO haccp_metingen 
                            (punt_id, chef_id, waarde, opmerking, actie_ondernomen)
                            VALUES (%s, %s, %s, %s, %s)
                        """, (punt['punt_id'], session['chef_id'], waarde, opmerking, actie))
                conn.commit()
                flash("HACCP metingen opgeslagen!", "success")
                return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))

        except Exception as e:
            conn.rollback()
            logger.error(f'Error processing HACCP checklist: {str(e)}')
            flash("Er is een fout opgetreden bij het verwerken van de checklist.", "danger")
        finally:
            cur.close()
            conn.close()

        return render_template('haccp/fill_checklist.html',
                            chef_naam=chef_naam,
                            checklist=checklist,
                            checkpunten=checkpunten)

    @app.route('/dashboard/<chef_naam>/haccp/reports', methods=['GET', 'POST'])
    @login_required
    def haccp_reports(chef_naam):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal alle checklists op voor het filter
            cur.execute("""
                SELECT * FROM haccp_checklists 
                WHERE chef_id = %s
            """, (session['chef_id'],))
            checklists = cur.fetchall()

            # Filter parameters met veilige defaults
            try:
                start_date = datetime.strptime(
                    request.args.get('start_date', ''), 
                    '%Y-%m-%d'
                ).strftime('%Y-%m-%d')
            except (ValueError, TypeError):
                start_date = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')

            try:
                end_date = datetime.strptime(
                    request.args.get('end_date', ''), 
                    '%Y-%m-%d'
                ).strftime('%Y-%m-%d')
            except (ValueError, TypeError):
                end_date = datetime.now().strftime('%Y-%m-%d')

            selected_checklist = request.args.get('checklist_id')

            # Query voor metingen
            query = """
                SELECT m.*, c.omschrijving, c.grenswaarde, cl.naam as checklist_naam
                FROM haccp_metingen m
                JOIN haccp_checkpunten c ON m.punt_id = c.punt_id
                JOIN haccp_checklists cl ON c.checklist_id = cl.checklist_id
                WHERE m.chef_id = %s
                AND DATE(m.timestamp) BETWEEN %s AND %s
            """
            params = [session['chef_id'], start_date, end_date]

            if selected_checklist:
                query += " AND cl.checklist_id = %s"
                params.append(selected_checklist)

            query += " ORDER BY m.timestamp DESC"
            cur.execute(query, tuple(params))
            metingen = cur.fetchall()

            # Bereken statistieken met veilige type conversies
            if metingen:
                afwijkingen = 0
                for m in metingen:
                    try:
                        waarde = float(m['waarde'] or 0)
                        grenswaarde = float(m['grenswaarde'] or 0)
                        if waarde > grenswaarde:
                            afwijkingen += 1
                    except (ValueError, TypeError):
                        continue
                compliance = ((len(metingen) - afwijkingen) / len(metingen) * 100) if metingen else 100
            else:
                compliance = 100

            if request.args.get('export'):
                return export_haccp_report(metingen, start_date, end_date, compliance)

            return render_template('haccp/reports.html',
                                chef_naam=chef_naam,
                                checklists=checklists,
                                metingen=metingen,
                                start_date=start_date,
                                end_date=end_date,
                                selected_checklist=selected_checklist,
                                compliance=compliance)

        except Exception as e:
            logger.error(f'Error in HACCP reports: {str(e)}')
            flash("Er is een fout opgetreden bij het ophalen van de HACCP rapportage.", "danger")
            return redirect(url_for('haccp_dashboard', chef_naam=chef_naam))
        finally:
            cur.close()
            conn.close()

    def export_haccp_report(metingen, start_date, end_date, compliance):
        doc = Document()
        doc.add_heading('HACCP Rapport', 0)
        
        # Voeg periode toe
        doc.add_paragraph(f'Periode: {start_date} t/m {end_date}')
        doc.add_paragraph(f'Compliance score: {compliance:.1f}%')

        # Maak tabel met metingen
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Datum'
        header_cells[1].text = 'Checklist'
        header_cells[2].text = 'Controlepunt'
        header_cells[3].text = 'Waarde'
        header_cells[4].text = 'Status'
        header_cells[5].text = 'Actie'

        # Voeg metingen toe
        for meting in metingen:
            row_cells = table.add_row().cells
            row_cells[0].text = meting['timestamp'].strftime('%d-%m-%Y %H:%M')
            row_cells[1].text = meting['checklist_naam']
            row_cells[2].text = meting['omschrijving']
            row_cells[3].text = f"{meting['waarde']}"
            
            if float(meting['waarde']) <= float(meting['grenswaarde']):
                row_cells[4].text = '✓ OK'
            else:
                row_cells[4].text = '⚠ Afwijking'
            
        header_cells[5].text = 'Actie'

        # Voeg metingen toe
        for meting in metingen:
            row_cells = table.add_row().cells
            row_cells[0].text = meting['timestamp'].strftime('%d-%m-%Y %H:%M')
            row_cells[1].text = meting['checklist_naam']
            row_cells[2].text = meting['omschrijving']
            row_cells[3].text = f"{meting['waarde']}"
            
            if float(meting['waarde']) <= float(meting['grenswaarde']):
                row_cells[4].text = '✓ OK'
            else:
                row_cells[4].text = '⚠ Afwijking'
            
            row_cells[5].text = meting['actie_ondernomen'] or 'Geen actie'
        
        # Initialize buffer before use
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'haccp_rapport_{start_date}_{end_date}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    @app.route('/dashboard/<chef_naam>/haccp/meting/<int:meting_id>/update', methods=['POST'])
    @login_required
    def update_haccp_meting(chef_naam, meting_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        try:
            # Debug logging
            app.logger.info(f"Received update request for meting {meting_id}")
            app.logger.info(f"Form data: {request.form}")

            waarde = request.form.get('waarde')
            actie_ondernomen = request.form.get('actie_ondernomen')

            if waarde is None:
                return jsonify({'success': False, 'error': 'Waarde ontbreekt'}), 400

            conn = get_db_connection()
            if conn is None:
                return jsonify({'success': False, 'error': 'Database verbinding mislukt'}), 500

            cursor = conn.cursor(dictionary=True)

            try:
                # Controleer eerst of de meting bestaat en van deze chef is
                cursor.execute("""
                    SELECT m.meting_id, c.grenswaarde 
                    FROM haccp_metingen m
                    JOIN haccp_checkpunten c ON m.punt_id = c.punt_id
                    WHERE m.meting_id = %s AND m.chef_id = %s
                """, (meting_id, session['chef_id']))
                meting = cursor.fetchone()
                if not meting:
                    return jsonify({'success': False, 'error': 'Meting niet gevonden'}), 404

                try:
                    waarde_float = float(waarde)
                except ValueError:
                    return jsonify({'success': False, 'error': 'Ongeldige waarde'}), 400

                # Update de meting (zonder explicit updated_at)
                cursor.execute("""
                    UPDATE haccp_metingen 
                    SET waarde = %s, actie_ondernomen = %s
                    WHERE meting_id = %s 
                    AND chef_id = %s
                """, (waarde_float, actie_ondernomen, meting_id, session['chef_id']))
                
                if cursor.rowcount == 0:
                    conn.rollback()
                    return jsonify({'success': False, 'error': 'Geen wijzigingen aangebracht'}), 400

                conn.commit()
                return jsonify({
                    'success': True,
                    'message': 'Meting succesvol bijgewerkt'
                })

            except mysql.connector.Error as e:
                conn.rollback()
                app.logger.error(f'Database error: {str(e)}')
                return jsonify({'success': False, 'error': f'Database fout: {str(e)}'}), 500
            finally:
                cursor.close()
                conn.close()

        except Exception as e:
            app.logger.error(f'Error updating HACCP measurement: {str(e)}')
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/dashboard/<chef_naam>/haccp/checklist/<int:checklist_id>/delete', methods=['POST'])
    @login_required
    def delete_haccp_checklist(chef_naam, checklist_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        try:
            conn = get_db_connection()
            if conn is None:
                return jsonify({'success': False, 'error': 'Database verbinding mislukt'}), 500
            cur = conn.cursor()

            # Controleer eerst of de checklist van deze chef is
            cur.execute("""
                SELECT chef_id FROM haccp_checklists 
                WHERE checklist_id = %s
            """, (checklist_id,))
            
            result = cur.fetchone()
            if not result or result[0] != session['chef_id']:
                return jsonify({'success': False, 'error': 'Checklist niet gevonden of geen toegang'}), 404

            # Verwijder eerst alle metingen van deze checklist
            cur.execute("""
                DELETE m FROM haccp_metingen m
                JOIN haccp_checkpunten c ON m.punt_id = c.punt_id
                WHERE c.checklist_id = %s
            """, (checklist_id,))

            # Verwijder dan alle checkpunten
            cur.execute("""
                DELETE FROM haccp_checkpunten 
                WHERE checklist_id = %s
            """, (checklist_id,))

            # Verwijder tenslotte de checklist zelf
            cur.execute("""
                DELETE FROM haccp_checklists 
                WHERE checklist_id = %s AND chef_id = %s
            """, (checklist_id, session['chef_id']))

            conn.commit()
            cur.close()
            conn.close()
            return jsonify({'success': True})

        except Exception as e:
            logger.error(f'Error deleting HACCP checklist: {str(e)}')
            return jsonify({'success': False, 'error': str(e)}), 500

    # -----------------------------------------------------------
    #  Terms and Conditions
    # -----------------------------------------------------------
    @app.route('/terms')
    def terms():
        return render_template('terms.html', form=FlaskForm())

    # Add alias for backward compatibility of 'quickstart' endpoint:
    @app.route('/quickstart', endpoint='quickstart')
    def quickstart_alias():
        return redirect(url_for('quickstart.quickstart_index'))

    def static_files(filename):
        return send_from_directory('static', filename)

    @app.route('/favicon.ico')
    def favicon():
        return send_from_directory(os.path.join(app.root_path, 'static', 'images'),
                                'favicon.png',
                                mimetype='image/png')

    @app.context_processor
    def inject_csrf_token():
        return dict(csrf_token=generate_csrf)  # Retourneer de functie zodat templates de token kunnen ophalen via csrf_token()

    # -----------------------------------------------------------
    #  Beheer Module (Stamgegevens)
    # -----------------------------------------------------------
    @app.route('/dashboard/<chef_naam>/beheer', methods=['GET', 'POST'])
    @login_required
    def beheer(chef_naam):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        leverancier_form = LeverancierForm(prefix="leverancier")
        eenheid_form = EenheidForm(prefix="eenheid")
        categorie_form = CategorieForm(prefix="categorie")
        dish_category_form = DishCategoryForm(prefix="dish_category")

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal bestaande stamgegevens op
            cur.execute("""
                SELECT * FROM leveranciers 
                WHERE chef_id = %s
                ORDER BY naam
            """, (session['chef_id'],))
            leveranciers = cur.fetchall()

            cur.execute("""
                SELECT * FROM eenheden 
                WHERE chef_id = %s
                ORDER BY naam
            """, (session['chef_id'],))
            eenheden = cur.fetchall()

            cur.execute("""
                SELECT * FROM categorieen 
                WHERE chef_id = %s
                ORDER BY naam
            """, (session['chef_id'],))
            categorieen = cur.fetchall()

            # Voeg dish_categories query toe
            cur.execute("""
                SELECT * FROM dish_categories 
                WHERE chef_id = %s
                ORDER BY volgorde, naam
            """, (session['chef_id'],))
            dish_categories = cur.fetchall()

            if request.method == 'POST':
                if leverancier_form.validate_on_submit() and leverancier_form.submit.data:
                    naam = leverancier_form.leverancier_naam.data
                    contact = leverancier_form.leverancier_contact.data
                    telefoon = leverancier_form.leverancier_telefoon.data
                    email = leverancier_form.leverancier_email.data
                    cur.execute("""
                        INSERT INTO leveranciers (chef_id, naam, contact, telefoon, email) 
                        VALUES (%s, %s, %s, %s, %s)
                    """, (session['chef_id'], naam, contact, telefoon, email))
                    conn.commit()
                    flash("Leverancier toegevoegd!", "success")
                    return redirect(url_for('beheer', chef_naam=chef_naam))

                elif eenheid_form.validate_on_submit() and eenheid_form.submit.data:
                    naam = eenheid_form.nieuwe_eenheid.data  # Changed from eenheid_nieuwe_eenheid
                    cur.execute("""
                        INSERT INTO eenheden (chef_id, naam)
                        VALUES (%s, %s)
                    """, (session['chef_id'], naam))
                    conn.commit()
                    flash("Eenheid toegevoegd!", "success")
                    return redirect(url_for('beheer', chef_naam=chef_naam))

                elif categorie_form.validate_on_submit() and categorie_form.submit.data:
                    naam = categorie_form.nieuwe_categorie.data  # Changed from categorie_nieuwe_categorie
                    cur.execute("""
                        INSERT INTO categorieen (chef_id, naam)
                        VALUES (%s, %s)
                    """, (session['chef_id'], naam))
                    conn.commit()
                    flash("Categorie toegevoegd!", "success")
                    return redirect(url_for('beheer', chef_naam=chef_naam))

                elif dish_category_form.validate_on_submit() and dish_category_form.submit.data:
                    naam = dish_category_form.nieuwe_dish_category.data  # Changed from dish_category_nieuwe_dish_category
                    volgorde = dish_category_form.volgorde.data  # Changed from dish_category_volgorde
                    cur.execute("""
                        INSERT INTO dish_categories (chef_id, naam, volgorde)
                        VALUES (%s, %s, %s)
                    """, (session['chef_id'], naam, volgorde))
                    conn.commit()
                    flash("Gang toegevoegd!", "success")
                    return redirect(url_for('beheer', chef_naam=chef_naam))

            return render_template('beheer.html',
                                chef_naam=chef_naam,
                                leveranciers=leveranciers,
                                eenheden=eenheden,
                                categorieen=categorieen,
                                dish_categories=dish_categories,
                                leverancier_form=leverancier_form,
                                eenheid_form=eenheid_form,
                                categorie_form=categorie_form,
                                dish_category_form=dish_category_form,
                                form=FlaskForm())

        except Exception as e:
            conn.rollback()
            flash(f"Fout: {str(e)}", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/dish_category/<int:category_id>/update', methods=['POST'])
    @login_required
    def update_dish_category(chef_naam, category_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = conn.cursor()
        try:
            naam = request.json.get('naam')
            volgorde = request.json.get('volgorde')

            cur.execute("""
                UPDATE dish_categories 
                SET naam = %s, volgorde = %s
                WHERE category_id = %s AND chef_id = %s
            """, (naam, volgorde, category_id, session['chef_id']))
            
            conn.commit()
            return jsonify({'success': True})
        except Exception as e:
            conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/dish_category/<int:category_id>/delete', methods=['POST'])
    @login_required
    def delete_dish_category(chef_naam, category_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = conn.cursor()
        try:
            # Controleer eerst of er gerechten zijn met deze categorie
            cur.execute("""
                SELECT COUNT(*) FROM dishes 
                WHERE chef_id = %s AND categorie = (
                    SELECT naam FROM dish_categories 
                    WHERE category_id = %s AND chef_id = %s
                )
            """, (session['chef_id'], category_id, session['chef_id']))
            
            if cur.fetchone()[0] > 0:
                return jsonify({
                    'success': False, 
                    'error': 'Deze categorie bevat nog gerechten en kan niet worden verwijderd.'
                }), 400

            cur.execute("""
                DELETE FROM dish_categories 
                WHERE category_id = %s AND chef_id = %s
            """, (category_id, session['chef_id']))
            
            conn.commit()
            return jsonify({'success': True})
        except Exception as e:
            conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/eenheid/<int:eenheid_id>/update', methods=['POST'])
    @login_required
    def update_eenheid(chef_naam, eenheid_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = conn.cursor()
        try:
            naam = request.json.get('naam')

            # Check if the unit exists and belongs to this chef
            cur.execute("""
                SELECT eenheid_id FROM eenheden 
                WHERE eenheid_id = %s AND chef_id = %s
            """, (eenheid_id, session['chef_id']))
            
            if not cur.fetchone():
                return jsonify({'success': False, 'error': 'Eenheid niet gevonden'}), 404

            # Update the unit
            cur.execute("""
                UPDATE eenheden 
                SET naam = %s
                WHERE eenheid_id = %s AND chef_id = %s
            """, (naam, eenheid_id, session['chef_id']))
            
            conn.commit()
            return jsonify({'success': True})
        except Exception as e:
            conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/eenheid/<int:eenheid_id>/delete', methods=['POST'])
    @login_required
    def delete_eenheid(chef_naam, eenheid_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = conn.cursor()
        try:
            # Check if any ingredients use this unit
            cur.execute("""
                SELECT COUNT(*) FROM ingredients 
                WHERE chef_id = %s AND eenheid = (
                    SELECT naam FROM eenheden 
                    WHERE eenheid_id = %s AND chef_id = %s
                )
            """, (session['chef_id'], eenheid_id, session['chef_id']))
            
            if cur.fetchone()[0] > 0:
                return jsonify({
                    'success': False, 
                    'error': 'Deze eenheid wordt nog gebruikt door ingrediënten en kan niet worden verwijderd.'
                }), 400

            cur.execute("""
                DELETE FROM eenheden 
                WHERE eenheid_id = %s AND chef_id = %s
            """, (eenheid_id, session['chef_id']))
            
            conn.commit()
            return jsonify({'success': True})
        except Exception as e:
            conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/categorie/<int:categorie_id>/update', methods=['POST'])
    @login_required
    def update_categorie(chef_naam, categorie_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        # Log alle inkomende data
        logger.info(f"Received update request for categorie {categorie_id}")
        logger.info(f"Request JSON: {request.get_json()}")
        logger.info(f"Headers: {dict(request.headers)}")
        
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            logger.warning(f"Authorization failed for chef {chef_naam}")
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            logger.error("Database connection failed")
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = None
        try:
            cur = conn.cursor(dictionary=True)
            
            # Valideer JSON data
            json_data = request.get_json()
            if not json_data:
                logger.error("No JSON data received")
                return jsonify({'success': False, 'error': 'Geen data ontvangen'}), 400

            naam = json_data.get('naam')
            if not naam:
                logger.error("Missing required field: naam")
                return jsonify({'success': False, 'error': 'Naam is verplicht'}), 400
                
            # Log query parameters
            logger.info(f"Executing query with params: categorie_id={categorie_id}, chef_id={session['chef_id']}, naam={naam}")
            cur.execute("""
                UPDATE categorieen 
                SET naam = %s
                WHERE categorie_id = %s AND chef_id = %s
            """, (naam, categorie_id, session['chef_id']))
            
            affected_rows = cur.rowcount
            logger.info(f"Query affected {affected_rows} rows")
            
            conn.commit()
            return jsonify({'success': True, 'message': 'Categorie succesvol bijgewerkt'})

        except Exception as e:
            logger.exception(f"Error in update_categorie: {str(e)}")
            if conn:
                conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            if cur:
                cur.close()
            if conn:
                conn.close()

    @app.route('/dashboard/<chef_naam>/categorie/<int:categorie_id>/delete', methods=['POST'])
    @login_required
    def delete_categorie(chef_naam, categorie_id):
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = conn.cursor()
        try:
            # Check if any ingredients use this category
            cur.execute("""
                SELECT COUNT(*) FROM ingredients 
                WHERE chef_id = %s AND categorie = (
                    SELECT naam FROM categorieen 
                    WHERE categorie_id = %s AND chef_id = %s
                )
            """, (session['chef_id'], categorie_id, session['chef_id']))
            
            if cur.fetchone()[0] > 0:
                return jsonify({
                    'success': False, 
                    'error': 'Deze categorie wordt nog gebruikt door ingrediënten en kan niet worden verwijderd.'
                }), 400

            cur.execute("""
                DELETE FROM categorieen 
                WHERE categorie_id = %s AND chef_id = %s
            """, (categorie_id, session['chef_id']))
            
            conn.commit()
            return jsonify({'success': True})
        except Exception as e:
            conn.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>')
    @login_required
    def dashboard(chef_naam):
        """Dashboard page after login"""
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))
            
        # Fetch counts for dishes and ingredients
        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('home'))
        
        try:
            cur = conn.cursor(dictionary=True)
            
            # Count dishes
            cur.execute("SELECT COUNT(*) as dish_count FROM dishes WHERE chef_id = %s", 
                        (session['chef_id'],))
            dish_count = cur.fetchone()['dish_count']
            
            # Count ingredients
            cur.execute("SELECT COUNT(*) as ingredient_count FROM ingredients WHERE chef_id = %s", 
                        (session['chef_id'],))
            ingredient_count = cur.fetchone()['ingredient_count']
            
            # Count suppliers
            cur.execute("SELECT COUNT(*) as supplier_count FROM leveranciers WHERE chef_id = %s", 
                        (session['chef_id'],))
            supplier_count = cur.fetchone()['supplier_count']
            
            cur.close()
            conn.close()
            
            return render_template('dashboard.html', 
                                 chef_naam=chef_naam, 
                                 dish_count=dish_count, 
                                 ingredient_count=ingredient_count, 
                                 supplier_count=supplier_count,
                                 form=FlaskForm())
        
        except Exception as e:
            if conn and conn.is_connected():
                cur.close()
                conn.close()
            logger.error(f"Error fetching dashboard counts: {str(e)}")
            flash("Er is een fout opgetreden bij het laden van uw gegevens.", "danger")
            return render_template('dashboard.html', chef_naam=chef_naam, form=FlaskForm())

    @app.route('/dashboard/<chef_naam>/dish/<int:dish_id>/costs', methods=['GET', 'POST'])
    @login_required
    def manage_dish_costs(chef_naam, dish_id):
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        
        form = FlaskForm()
        if 'chef_id' not in session or session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('all_dishes'))
        
        cur = conn.cursor(dictionary=True)

        try:
            # Fetch dish info
            cur.execute("""
                SELECT * FROM dishes 
                WHERE dish_id = %s AND chef_id = %s
            """, (dish_id, session['chef_id']))
            gerecht = cur.fetchone()

            if not gerecht:
                flash("Gerecht niet gevonden.", "danger")
                return redirect(url_for('all_dishes'))

            if request.method == 'POST':
                if not form.validate_on_submit():
                    flash("Ongeldige CSRF-token.", "danger")
                    return redirect(url_for('manage_dish_costs', chef_naam=chef_naam, dish_id=dish_id))

                if 'ingredientForm' in request.form:
                    ingredient_id = request.form.get('ingredient_id')
                    hoeveelheid = request.form.get('hoeveelheid')

                    # Calculate total price
                    cur.execute("""
                        SELECT prijs_per_eenheid 
                        FROM ingredients 
                        WHERE ingredient_id = %s 
                        AND chef_id = %s
                    """, (ingredient_id, session['chef_id']))
                    ingredient_info = cur.fetchone()
                    
                    if ingredient_info:
                        prijs_per_eenheid = float(ingredient_info['prijs_per_eenheid'])
                        prijs_totaal = round(float(hoeveelheid) * prijs_per_eenheid, 5)

                        try:
                            # Check if ingredient already exists
                            cur.execute("""
                                SELECT * FROM dish_ingredients 
                                WHERE dish_id = %s AND ingredient_id = %s
                            """, (dish_id, ingredient_id))
                            if cur.fetchone():
                                flash("Dit ingrediënt is al gekoppeld aan de kostprijsberekening.", "danger")
                            else:
                                cur.execute("""
                                    INSERT INTO dish_ingredients 
                                    (dish_id, ingredient_id, hoeveelheid, prijs_totaal)
                                    VALUES (%s, %s, %s, %s)
                                """, (dish_id, ingredient_id, hoeveelheid, prijs_totaal))
                                conn.commit()
                                flash("Ingrediënt toegevoegd!", "success")
                        except Exception as e:
                            conn.rollback()
                            flash(f"Fout bij toevoegen: {str(e)}", "danger")

            # Get all ingredients
            cur.execute("""
                SELECT * FROM ingredients 
                WHERE chef_id = %s
                ORDER BY naam
            """, (session['chef_id'],))
            alle_ingredienten = cur.fetchall()

            # Get linked ingredients
            cur.execute("""
                SELECT di.*, i.naam AS ingredient_naam, i.eenheid, i.prijs_per_eenheid
                FROM dish_ingredients di
                JOIN ingredients i ON di.ingredient_id = i.ingredient_id
                WHERE di.dish_id = %s
            """, (dish_id,))
            gerecht_ingredienten = cur.fetchall()

            # Calculate total cost
            totaal_ingredient_prijs = sum(float(gi['prijs_totaal']) for gi in gerecht_ingredienten)

            return render_template(
                'manage_dish_costs.html',
                chef_naam=chef_naam,
                gerecht=gerecht,
                alle_ingredienten=alle_ingredienten,
                gerecht_ingredienten=gerecht_ingredienten,
                totaal_ingredient_prijs=totaal_ingredient_prijs,
                form=form
            )

        except Exception as e:
            logger.error(f'Error in manage_dish_costs: {str(e)}')
            flash("Er is een fout opgetreden.", "danger")
            return redirect(url_for('all_dishes'))
        finally:
            cur.close()
            conn.close()

    @app.route('/print_menu')
    @login_required
    def print_menu():
        """
        Pagina om menukaart samen te stellen en te printen.
        """
        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('home'))
        cur = conn.cursor(dictionary=True)

        try:
            # Haal alle gerechten van de ingelogde chef op
            cur.execute("""
                SELECT d.*, c.naam as chef_naam 
                FROM dishes d
                JOIN chefs c ON d.chef_id = c.chef_id
                WHERE d.chef_id = %s
                ORDER BY d.categorie, d.naam
            """, (session['chef_id'],))
            gerechten = cur.fetchall()

            return render_template('print_menu.html', 
                                gerechten=gerechten,
                                form=FlaskForm())
        except Exception as e:
            flash(f"Error: {str(e)}", "danger")
            return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
        finally:
            cur.close()
            conn.close()

    @app.route('/dashboard/<chef_naam>/dishes/<int:dish_id>/price', methods=['POST'])
    @login_required
    def update_dish_price(chef_naam, dish_id):
        """Update verkoopprijs van een gerecht"""
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        
        form = FlaskForm()
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token.", "danger")
            return redirect(url_for('all_dishes'))

        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('all_dishes'))
        
        cur = conn.cursor()

        try:
            # Handle both dot and comma as decimal separator
            prijs = request.form.get('verkoopprijs', '0.00')
            prijs = prijs.replace(',', '.')  # Convert comma to dot
            
            try:
                # Validate and convert to proper decimal format
                prijs_float = float(prijs)
                if prijs_float < 0:
                    raise ValueError("Prijs mag niet negatief zijn")
                
                nieuwe_verkoopprijs = '{:.2f}'.format(prijs_float)
                
                cur.execute("""
                    UPDATE dishes
                    SET verkoopprijs = %s
                    WHERE dish_id = %s AND chef_id = %s
                """, (nieuwe_verkoopprijs, dish_id, session['chef_id']))
                
                conn.commit()
                flash("Verkoopprijs bijgewerkt!", "success")
            except ValueError as e:
                flash(f"Ongeldige prijs ingevoerd: {str(e)}", "danger")
                return redirect(url_for('all_dishes'))
                
        except Exception as e:
            conn.rollback()
            flash(f"Fout bij bijwerken prijs: {str(e)}", "danger")
        finally:
            cur.close()
            conn.close()

        return redirect(url_for('all_dishes'))

    @app.route('/dashboard/<chef_naam>/recalculate-prices', methods=['POST'])
    @login_required
    def recalculate_all_prices(chef_naam):
        """Handmatig alle kostprijzen herberekenen"""
        if session['chef_naam'] != chef_naam:
            return jsonify({'success': False, 'error': 'Geen toegang'}), 403

        conn = get_db_connection()
        if conn is None:
            return jsonify({'success': False, 'error': 'Database verbindingsfout'}), 500

        cur = conn.cursor(dictionary=True)
        try:
            # Haal alle gerechten op met hun ingrediënten
            cur.execute("""
                SELECT d.dish_id, di.ingredient_id, di.hoeveelheid, i.prijs_per_eenheid
                FROM dishes d
                JOIN dish_ingredients di ON d.dish_id = di.dish_id
                JOIN ingredients i ON di.ingredient_id = i.ingredient_id
                WHERE d.chef_id = %s
            """, (session['chef_id'],))
            
            all_ingredients = cur.fetchall()
            
            # Groepeer per gerecht
            dishes = {}
            for row in all_ingredients:
                if row['dish_id'] not in dishes:
                    dishes[row['dish_id']] = []
                dishes[row['dish_id']].append(row)
            
            # Update kostprijs per gerecht
            updated_dishes = 0
            for dish_id, ingredients in dishes.items():
                total_cost = sum(ing['hoeveelheid'] * float(ing['prijs_per_eenheid']) for ing in ingredients)
                
                # Update dish_ingredients tabel
                for ingredient in ingredients:
                    price = ingredient['hoeveelheid'] * float(ingredient['prijs_per_eenheid'])
                    cur.execute("""
                        UPDATE dish_ingredients 
                        SET prijs_totaal = %s
                        WHERE dish_id = %s AND ingredient_id = %s
                    """, (price, dish_id, ingredient['ingredient_id']))
                
                updated_dishes += 1

            conn.commit()
            return jsonify({'success': True, 'updated_dishes': updated_dishes})

        except Exception as e:
            conn.rollback()
            logger.error(f"Error recalculating prices: {str(e)}")
            return jsonify({'success': False, 'error': str(e)}), 500
        finally:
            cur.close()
            conn.close()

    @app.after_request
    def after_request(response):
        """Clear skipped ingredients from session after displaying them once"""
        if request.endpoint == 'manage_suppliers' and 'skipped_ingredients' in session:
            # Keep skipped ingredients for one more request after redirect
            # This enables the modal to show them after page refresh
            if request.method == 'GET':
                # Only clear after GET request (after redirect from POST)
                session.pop('skipped_ingredients', None)
        return response

    # -----------------------------------------------------------
    #  Backup & Restore Functionality
    # -----------------------------------------------------------
    @app.route('/backup/<chef_naam>', methods=['GET'])
    @login_required
    def backup_chef_data(chef_naam):
        """Generate a complete backup of chef data"""
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
        
        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))
            
        try:
            chef_id = session['chef_id']
            
            conn = get_db_connection()
            if conn is None:
                flash("Database connection error.", "danger")
                return redirect(url_for('dashboard', chef_naam=chef_naam))
                
            cur = conn.cursor(dictionary=True)
            
            # Create a backup dictionary
            backup_data = {
                'metadata': {
                    'chef_id': chef_id,
                    'chef_naam': chef_naam,
                    'created_at': datetime.now().isoformat(),
                    'version': '1.0'
                },
                'data': {}
            }
            
            # Backup dishes
            cur.execute("SELECT * FROM dishes WHERE chef_id = %s", (chef_id,))
            backup_data['data']['dishes'] = cur.fetchall()
            
            # Backup dish_ingredients
            cur.execute("""
                SELECT di.* 
                FROM dish_ingredients di
                JOIN dishes d ON di.dish_id = d.dish_id
                WHERE d.chef_id = %s
            """, (chef_id,))
            backup_data['data']['dish_ingredients'] = cur.fetchall()
            
            # Backup dish_allergenen
            cur.execute("""
                SELECT da.* 
                FROM dish_allergenen da
                JOIN dishes d ON da.dish_id = d.dish_id
                WHERE d.chef_id = %s
            """, (chef_id,))
            backup_data['data']['dish_allergenen'] = cur.fetchall()
            
            # Backup dish_dieten
            cur.execute("""
                SELECT dd.* 
                FROM dish_dieten dd
                JOIN dishes d ON dd.dish_id = d.dish_id
                WHERE d.chef_id = %s
            """, (chef_id,))
            backup_data['data']['dish_dieten'] = cur.fetchall()
            
            # Backup ingredients
            cur.execute("SELECT * FROM ingredients WHERE chef_id = %s", (chef_id,))
            backup_data['data']['ingredients'] = cur.fetchall()
            
            # Backup leveranciers
            cur.execute("SELECT * FROM leveranciers WHERE chef_id = %s", (chef_id,))
            backup_data['data']['leveranciers'] = cur.fetchall()
            
            # Backup eenheden
            cur.execute("SELECT * FROM eenheden WHERE chef_id = %s", (chef_id,))
            backup_data['data']['eenheden'] = cur.fetchall()
            
            # Backup categorieen
            cur.execute("SELECT * FROM categorieen WHERE chef_id = %s", (chef_id,))
            backup_data['data']['categorieen'] = cur.fetchall()
            
            # Backup dish_categories
            cur.execute("SELECT * FROM dish_categories WHERE chef_id = %s", (chef_id,))
            backup_data['data']['dish_categories'] = cur.fetchall()
            
            # Backup HACCP data if exists
            cur.execute("SELECT * FROM haccp_checklists WHERE chef_id = %s", (chef_id,))
            backup_data['data']['haccp_checklists'] = cur.fetchall()
            
            # Get checkpoints for these checklists
            haccp_ids = [row['checklist_id'] for row in backup_data['data']['haccp_checklists']]
            if haccp_ids:
                placeholder = ','.join(['%s'] * len(haccp_ids))
                cur.execute(f"""
                    SELECT * FROM haccp_checkpunten 
                    WHERE checklist_id IN ({placeholder})
                """, tuple(haccp_ids))
                backup_data['data']['haccp_checkpunten'] = cur.fetchall()
                
                # Get measurements for these chef's checkpoints
                cur.execute("""
                    SELECT m.* 
                    FROM haccp_metingen m
                    JOIN haccp_checkpunten p ON m.punt_id = p.punt_id
                    JOIN haccp_checklists c ON p.checklist_id = c.checklist_id
                    WHERE c.chef_id = %s
                """, (chef_id,))
                backup_data['data']['haccp_metingen'] = cur.fetchall()
            
            # Backup takenboek if exists
            try:
                cur.execute("SELECT * FROM tasks WHERE chef_id = %s", (chef_id,))
                backup_data['data']['tasks'] = cur.fetchall()
            except:
                # Table might not exist
                pass
                
            cur.close()
            conn.close()
            
            # Convert Decimal objects to strings
            def decimal_serializer(obj):
                if isinstance(obj, Decimal):
                    return str(obj)
                # Convert datetime objects
                if isinstance(obj, datetime):
                    return obj.isoformat()
                raise TypeError(f"Type {type(obj)} not serializable")
                
            # Generate JSON file
            now = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_filename = f"echef_backup_{chef_naam}_{now}.json"
            
            # Return the file as a download
            response = make_response(json.dumps(backup_data, default=decimal_serializer, indent=2))
            response.headers.set('Content-Type', 'application/json')
            response.headers.set('Content-Disposition', f'attachment; filename={backup_filename}')
            return response
            
        except Exception as e:
            logger.error(f"Error creating backup: {str(e)}", exc_info=True)
            flash(f"Er is een fout opgetreden bij het maken van de backup: {str(e)}", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))

    @app.route('/restore/<chef_naam>', methods=['POST'])
    @login_required
    def restore_chef_data(chef_naam):
        """Restore chef data from backup file"""
        if session['chef_naam'] != chef_naam:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('auth.login'))
            
        if 'chef_id' not in session:
            flash("Geen toegang. Log opnieuw in.", "danger")
            return redirect(url_for('login'))
            
        form = FlaskForm()
        if not form.validate_on_submit():
            flash("Ongeldige CSRF-token.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
            
        if 'backup_file' not in request.files:
            flash("Geen backup bestand geselecteerd.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
            
        file = request.files['backup_file']
        if file.filename == '':
            flash("Geen backup bestand geselecteerd.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
            
        try:
            # Verify it's a JSON file
            if not file.filename.endswith('.json'):
                flash("Ongeldig bestandsformaat. Upload een .json backup bestand.", "danger")
                return redirect(url_for('dashboard', chef_naam=chef_naam))
                
            # Read and validate the backup file
            backup_content = file.read().decode('utf-8')
            backup_data = json.loads(backup_content)
            
            # Basic validation
            if 'metadata' not in backup_data or 'data' not in backup_data:
                flash("Ongeldig backup formaat.", "danger")
                return redirect(url_for('dashboard', chef_naam=chef_naam))
                
            # Verify the backup belongs to this chef
            if str(backup_data['metadata']['chef_id']) != str(session['chef_id']):
                flash("Deze backup hoort niet bij jouw account.", "danger")
                return redirect(url_for('dashboard', chef_naam=chef_naam))
                
            # Start restoration process
            conn = get_db_connection()
            if conn is None:
                flash("Database connection error.", "danger")
                return redirect(url_for('dashboard', chef_naam=chef_naam))
                
            cur = conn.cursor(dictionary=True)
            chef_id = session['chef_id']
            
            # Confirm restoration
            confirm = request.form.get('confirm_restore', 'false')
            if confirm != 'true':
                flash("Om een backup te herstellen, moet je de waarschuwing bevestigen.", "warning")
                return redirect(url_for('dashboard', chef_naam=chef_naam))
                
            # Start transaction
            conn.start_transaction()
            
            try:
                # Clear existing data (in reverse order of dependencies)
                tables_to_clear = [
                    "haccp_metingen",
                    "haccp_checkpunten",
                    "haccp_checklists",
                    "dish_dieten",
                    "dish_allergenen",
                    "dish_ingredients",
                    "dishes",
                    "ingredients",
                    "leveranciers",
                    "eenheden",
                    "categorieen",
                    "dish_categories"
                ]
                
                for table in tables_to_clear:
                    # Some tables need special handling due to foreign key relationships
                    if table == "dish_ingredients":
                        cur.execute(f"""
                            DELETE di FROM dish_ingredients di 
                            JOIN dishes d ON di.dish_id = d.dish_id 
                            WHERE d.chef_id = %s
                        """, (chef_id,))
                    elif table == "dish_allergenen":
                        cur.execute(f"""
                            DELETE da FROM dish_allergenen da 
                            JOIN dishes d ON da.dish_id = d.dish_id 
                            WHERE d.chef_id = %s
                        """, (chef_id,))
                    elif table == "dish_dieten":
                        cur.execute(f"""
                            DELETE dd FROM dish_dieten dd 
                            JOIN dishes d ON dd.dish_id = d.dish_id 
                            WHERE d.chef_id = %s
                        """, (chef_id,))
                    elif table == "haccp_metingen":
                        # Skip if not in backup
                        if 'haccp_metingen' not in backup_data['data']:
                            continue
                        cur.execute(f"""
                            DELETE m FROM haccp_metingen m
                            JOIN haccp_checkpunten p ON m.punt_id = p.punt_id
                            JOIN haccp_checklists c ON p.checklist_id = c.checklist_id
                            WHERE c.chef_id = %s
                        """, (chef_id,))
                    elif table == "haccp_checkpunten":
                        # Skip if not in backup
                        if 'haccp_checkpunten' not in backup_data['data']:
                            continue
                        cur.execute(f"""
                            DELETE p FROM haccp_checkpunten p
                            JOIN haccp_checklists c ON p.checklist_id = c.checklist_id
                            WHERE c.chef_id = %s
                        """, (chef_id,))
                    else:
                        # Skip if not in backup
                        if table not in backup_data['data']:
                            continue
                        cur.execute(f"DELETE FROM {table} WHERE chef_id = %s", (chef_id,))
                
                # Restore each table (in order of dependencies)
                restore_order = [
                    "leveranciers", 
                    "categorieen", 
                    "eenheden", 
                    "dish_categories",
                    "ingredients",
                    "dishes",
                    "dish_ingredients",
                    "dish_allergenen",
                    "dish_dieten",
                    "haccp_checklists",
                    "haccp_checkpunten",
                    "haccp_metingen"
                ]
                
                # Add tasks if present
                if 'tasks' in backup_data['data']:
                    restore_order.append("tasks")
                
                for table in restore_order:
                    # Skip if not in backup
                    if table not in backup_data['data'] or not backup_data['data'][table]:
                        continue
                    
                    for row in backup_data['data'][table]:
                        # Convert keys to list for consistent ordering
                        columns = list(row.keys())
                        values = [row[col] for col in columns]
                        
                        # Build placeholders
                        placeholders = ', '.join(['%s'] * len(values))
                        columns_str = ', '.join(columns)
                        
                        # Insert query with ignore duplicate keys
                        query = f"INSERT IGNORE INTO {table} ({columns_str}) VALUES ({placeholders})"
                        cur.execute(query, tuple(values))
                
                conn.commit()
                flash("Backup succesvol hersteld! Je data is teruggezet naar de staat van de backup.", "success")
                
            except Exception as e:
                conn.rollback()
                logger.error(f"Error restoring backup: {str(e)}", exc_info=True)
                flash(f"Fout bij het herstellen van de backup: {str(e)}", "danger")
                
            finally:
                cur.close()
                conn.close()
                
            return redirect(url_for('dashboard', chef_naam=chef_naam))
            
        except json.JSONDecodeError:
            flash("Het geselecteerde bestand is geen geldige JSON backup.", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))
            
        except Exception as e:
            logger.error(f"Error processing backup file: {str(e)}", exc_info=True)
            flash(f"Fout bij het verwerken van het backup bestand: {str(e)}", "danger")
            return redirect(url_for('dashboard', chef_naam=chef_naam))

    return app

# Move configuration constants outside create_app
DB_NAME = os.getenv("DB_NAME")
DB_USER = os.getenv("DB_USER")
DB_PASSWORD = os.getenv("DB_PASSWORD")
DB_HOST = os.getenv("DB_HOST")
DB_PORT = os.getenv("DB_PORT")

# -----------------------------------------------------------
# Start de server alleen lokaal
# -----------------------------------------------------------
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    application = create_app()
    if os.environ.get('FLASK_ENV') == 'development':
        # Use waitress for local development
        from waitress import serve
        serve(application, host='0.0.0.0', port=port)
    else:
        # Use gunicorn in production (Heroku)
        application.run(host='0.0.0.0', port=port)