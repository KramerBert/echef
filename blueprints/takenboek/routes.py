import re
from flask import (
    render_template, redirect, url_for, flash, request, 
    session, jsonify, current_app, send_file
)
from . import bp
from utils.db import get_db_connection
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

# Set up logging
logger = logging.getLogger(__name__)

def login_required(f):
    """Decorator to ensure user is logged in"""
    def decorated_function(*args, **kwargs):
        if 'chef_id' not in session:
            flash("Je moet ingelogd zijn om deze pagina te bekijken.", "danger")
            return redirect(url_for('auth.login'))
        return f(*args, **kwargs)
    decorated_function.__name__ = f.__name__
    return decorated_function

@bp.route('/dashboard/<chef_naam>/takenboek', endpoint='index')
@login_required
def takenboek(chef_naam):
    """Display the student task book"""
    if session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('auth.login'))
    
    if 'chef_id' not in session or session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('login'))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('dashboard', chef_naam=chef_naam))
    
    cur = conn.cursor(dictionary=True)
    
    try:
        # Check if the chef already has tasks
        cur.execute("""
            SELECT COUNT(*) as count 
            FROM student_tasks 
            WHERE chef_id = %s
        """, (session['chef_id'],))
        
        task_count = cur.fetchone()['count']
        has_tasks = task_count > 0
        
        # Get all tasks if they exist
        tasks = []
        if has_tasks:
            cur.execute("""
                SELECT * FROM student_tasks
                WHERE chef_id = %s
                ORDER BY blok, onderdeel, techniek
            """, (session['chef_id'],))
            tasks = cur.fetchall()
        
        return render_template('takenboek/index.html', 
                              chef_naam=chef_naam, 
                              has_tasks=has_tasks, 
                              tasks=tasks,
                              form=dict(csrf_token=request.form.get('csrf_token', '')))
    
    except Exception as e:
        logger.error(f'Error in takenboek: {str(e)}')
        flash(f"Er is een fout opgetreden: {str(e)}", "danger")
        return redirect(url_for('dashboard', chef_naam=chef_naam))
    
    finally:
        cur.close()
        conn.close()

@bp.route('/dashboard/<chef_naam>/takenboek/create', methods=['POST'])
@login_required
def create_takenboek(chef_naam):
    """Create a new task book using the database trigger"""
    if session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('auth.login'))
    
    if 'chef_id' not in session or session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('login'))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('takenboek.index', chef_naam=chef_naam))
    
    cur = conn.cursor(dictionary=True)
    
    try:
        # Disable safe update mode temporarily
        cur.execute("SET SQL_SAFE_UPDATES = 0;")
        
        # Check if the chef already has tasks
        cur.execute("""
            SELECT COUNT(*) as count 
            FROM student_tasks 
            WHERE chef_id = %s
        """, (session['chef_id'],))
        
        task_count = cur.fetchone()['count']
        
        if task_count > 0:
            flash("Je hebt al een takenboek.", "warning")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
        
        # Get the default trigger template
        cur.execute("""
            SELECT trigger_sql 
            FROM task_trigger_template 
            WHERE is_default = TRUE 
            LIMIT 1
        """)
        
        template = cur.fetchone()
        
        if not template:
            flash("Er is geen standaard takenboek template gevonden.", "danger")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
        
        # Process the tasks individually to avoid SQL syntax issues
        trigger_sql = template['trigger_sql']
        chef_id = session['chef_id']
        
        # Extract individual task data using regex pattern matching instead of SQL parsing
        # Match pattern for each value tuple in the format (:chef_id, "val1", "val2", "val3", "val4")
        pattern = r'\(:chef_id,\s*"([^"]+)",\s*"([^"]+)",\s*"([^"]+)",\s*"((?:[^"\\]|\\.)+)"\)'
        matches = re.findall(pattern, trigger_sql)
        
        tasks_inserted = 0
        
        for match in matches:
            if len(match) == 4:
                blok = match[0]
                onderdeel = match[1]
                techniek = match[2]
                taak = match[3]
                
                cur.execute("""
                    INSERT INTO student_tasks (chef_id, blok, onderdeel, techniek, taak)
                    VALUES (%s, %s, %s, %s, %s)
                """, (chef_id, blok, onderdeel, techniek, taak))
                
                tasks_inserted += 1
        
        # Re-enable safe update mode
        cur.execute("SET SQL_SAFE_UPDATES = 1;")
        
        conn.commit()
        flash(f"Takenboek succesvol aangemaakt met {tasks_inserted} taken!", "success")
        
    except Exception as e:
        conn.rollback()
        logger.error(f'Error creating task book: {str(e)}')
        flash(f"Fout bij aanmaken takenboek: {str(e)}", "danger")
    
    finally:
        cur.close()
        conn.close()
    
    return redirect(url_for('takenboek.index', chef_naam=chef_naam))

@bp.route('/dashboard/<chef_naam>/takenboek/add', methods=['GET', 'POST'])
@login_required
def add_task(chef_naam):
    """Add a new task to the task book"""
    if session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('auth.login'))
    
    if 'chef_id' not in session or session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        blok = request.form.get('blok')
        onderdeel = request.form.get('onderdeel')
        techniek = request.form.get('techniek')
        taak = request.form.get('taak')
        
        # Validate form data
        if not all([blok, onderdeel, techniek, taak]):
            flash("Alle velden zijn verplicht.", "danger")
            return render_template('takenboek/add_task.html', 
                                chef_naam=chef_naam,
                                form=dict(csrf_token=request.form.get('csrf_token', '')))
        
        conn = get_db_connection()
        if conn is None:
            flash("Database connection error.", "danger")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
        
        cur = conn.cursor()
        
        try:
            cur.execute("""
                INSERT INTO student_tasks (chef_id, blok, onderdeel, techniek, taak)
                VALUES (%s, %s, %s, %s, %s)
            """, (session['chef_id'], blok, onderdeel, techniek, taak))
            
            conn.commit()
            flash("Taak succesvol toegevoegd!", "success")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
            
        except Exception as e:
            conn.rollback()
            logger.error(f'Error adding task: {str(e)}')
            flash(f"Fout bij toevoegen taak: {str(e)}", "danger")
        
        finally:
            cur.close()
            conn.close()
    
    return render_template('takenboek/add_task.html', 
                          chef_naam=chef_naam,
                          form=dict(csrf_token=request.form.get('csrf_token', '')))

@bp.route('/dashboard/<chef_naam>/takenboek/edit/<int:task_id>', methods=['GET', 'POST'])
@login_required
def edit_task(chef_naam, task_id):
    """Edit an existing task"""
    if session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('auth.login'))
    
    if 'chef_id' not in session or session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('login'))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('takenboek.index', chef_naam=chef_naam))
    
    cur = conn.cursor(dictionary=True)
    
    try:
        # Get the task
        cur.execute("""
            SELECT * 
            FROM student_tasks 
            WHERE task_id = %s AND chef_id = %s
        """, (task_id, session['chef_id']))
        
        task = cur.fetchone()
        
        if not task:
            flash("Taak niet gevonden.", "danger")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
        
        if request.method == 'POST':
            blok = request.form.get('blok')
            onderdeel = request.form.get('onderdeel')
            techniek = request.form.get('techniek')
            taak = request.form.get('taak')
            uitgevoerd_op = request.form.get('uitgevoerd_op')
            opmerkingen = request.form.get('opmerkingen')
            
            # Update the task
            cur.execute("""
                UPDATE student_tasks
                SET blok = %s, onderdeel = %s, techniek = %s, taak = %s, 
                    uitgevoerd_op = %s, opmerkingen = %s
                WHERE task_id = %s AND chef_id = %s
            """, (blok, onderdeel, techniek, taak, uitgevoerd_op or None, 
                  opmerkingen, task_id, session['chef_id']))
            
            conn.commit()
            flash("Taak succesvol bijgewerkt!", "success")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
        
        return render_template('takenboek/edit_task.html', 
                               chef_naam=chef_naam, 
                               task=task,
                               form=dict(csrf_token=request.form.get('csrf_token', '')))
    
    except Exception as e:
        if request.method == 'POST':
            conn.rollback()
            logger.error(f'Error updating task: {str(e)}')
            flash(f"Fout bij bijwerken taak: {str(e)}", "danger")
        
        logger.error(f'Error in edit_task: {str(e)}')
        flash(f"Er is een fout opgetreden: {str(e)}", "danger")
        return redirect(url_for('takenboek.index', chef_naam=chef_naam))
    
    finally:
        cur.close()
        conn.close()

@bp.route('/dashboard/<chef_naam>/takenboek/delete/<int:task_id>', methods=['POST'])
@login_required
def delete_task(chef_naam, task_id):
    """Delete a task from the task book"""
    if session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('auth.login'))
    
    if 'chef_id' not in session or session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('login'))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('takenboek.index', chef_naam=chef_naam))
    
    cur = conn.cursor()
    
    try:
        # Delete the task
        cur.execute("""
            DELETE FROM student_tasks 
            WHERE task_id = %s AND chef_id = %s
        """, (task_id, session['chef_id']))
        
        if cur.rowcount == 0:
            flash("Taak niet gevonden.", "danger")
        else:
            conn.commit()
            flash("Taak succesvol verwijderd!", "success")
    
    except Exception as e:
        conn.rollback()
        logger.error(f'Error deleting task: {str(e)}')
        flash(f"Fout bij verwijderen taak: {str(e)}", "danger")
    
    finally:
        cur.close()
        conn.close()
    
    return redirect(url_for('takenboek.index', chef_naam=chef_naam))

@bp.route('/admin/takenboek/trigger', methods=['GET', 'POST'])
@login_required
def admin_task_trigger():
    """Admin interface for modifying the task book trigger template"""
    # Check if user is admin
    if 'is_admin' not in session or not session['is_admin']:
        flash("Je hebt geen toegang tot deze pagina.", "danger")
        return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
    
    cur = conn.cursor(dictionary=True)
    
    try:
        if request.method == 'POST':
            naam = request.form.get('naam')
            omschrijving = request.form.get('omschrijving')
            trigger_sql = request.form.get('trigger_sql')
            is_default = 'is_default' in request.form
            
            # Insert new template
            cur.execute("""
                INSERT INTO task_trigger_template (naam, omschrijving, trigger_sql, is_default)
                VALUES (%s, %s, %s, %s)
            """, (naam, omschrijving, trigger_sql, is_default))
            
            # If this is the new default, unset others
            if is_default:
                cur.execute("""
                    UPDATE task_trigger_template
                    SET is_default = FALSE
                    WHERE template_id != LAST_INSERT_ID()
                """)
            
            conn.commit()
            flash("Template succesvol toegevoegd!", "success")
        
        # Get all templates
        cur.execute("SELECT * FROM task_trigger_template ORDER BY is_default DESC, created_at DESC")
        templates = cur.fetchall()
        
        return render_template('admin/task_trigger.html', 
                               templates=templates,
                               form=dict(csrf_token=request.form.get('csrf_token', '')))
    
    except Exception as e:
        if request.method == 'POST':
            conn.rollback()
            logger.error(f'Error adding trigger template: {str(e)}')
            flash(f"Fout bij toevoegen template: {str(e)}", "danger")
        
        logger.error(f'Error in admin_task_trigger: {str(e)}')
        flash(f"Er is een fout opgetreden: {str(e)}", "danger")
        return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
    
    finally:
        cur.close()
        conn.close()

@bp.route('/admin/takenboek/trigger/<int:template_id>/set-default', methods=['POST'])
@login_required
def set_default_trigger(template_id):
    """Set a template as the default trigger"""
    # Check if user is admin
    if 'is_admin' not in session or not session['is_admin']:
        flash("Je hebt geen toegang tot deze pagina.", "danger")
        return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('takenboek.admin_task_trigger'))
    
    cur = conn.cursor()
    
    try:
        # Reset all templates to not default
        cur.execute("UPDATE task_trigger_template SET is_default = FALSE")
        
        # Set the selected template as default
        cur.execute("""
            UPDATE task_trigger_template 
            SET is_default = TRUE 
            WHERE template_id = %s
        """, (template_id,))
        
        conn.commit()
        flash("Standaard trigger template bijgewerkt!", "success")
    
    except Exception as e:
        conn.rollback()
        logger.error(f'Error setting default trigger: {str(e)}')
        flash(f"Fout bij instellen standaard trigger: {str(e)}", "danger")
    
    finally:
        cur.close()
        conn.close()
    
    return redirect(url_for('takenboek.admin_task_trigger'))

@bp.route('/admin/takenboek/trigger/<int:template_id>/delete', methods=['POST'])
@login_required
def delete_trigger(template_id):
    """Delete a trigger template"""
    # Check if user is admin
    if 'is_admin' not in session or not session['is_admin']:
        flash("Je hebt geen toegang tot deze pagina.", "danger")
        return redirect(url_for('dashboard', chef_naam=session.get('chef_naam')))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('takenboek.admin_task_trigger'))
    
    cur = conn.cursor(dictionary=True)
    
    try:
        # Check if this is the default template
        cur.execute("""
            SELECT is_default 
            FROM task_trigger_template 
            WHERE template_id = %s
        """, (template_id,))
        
        template = cur.fetchone()
        
        if not template:
            flash("Template niet gevonden.", "danger")
            return redirect(url_for('takenboek.admin_task_trigger'))
        
        if template['is_default']:
            flash("Je kunt de standaard template niet verwijderen.", "danger")
            return redirect(url_for('takenboek.admin_task_trigger'))
        
        # Delete the template
        cur.execute("""
            DELETE FROM task_trigger_template 
            WHERE template_id = %s AND is_default = FALSE
        """, (template_id,))
        
        conn.commit()
        flash("Template succesvol verwijderd!", "success")
    
    except Exception as e:
        conn.rollback()
        logger.error(f'Error deleting trigger: {str(e)}')
        flash(f"Fout bij verwijderen template: {str(e)}", "danger")
    
    finally:
        cur.close()
        conn.close()
    
    return redirect(url_for('takenboek.admin_task_trigger'))

@bp.route('/dashboard/<chef_naam>/takenboek/export', methods=['POST'])
@login_required
def export_tasks(chef_naam):
    """Export tasks to Word document"""
    if session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('auth.login'))
    
    if 'chef_id' not in session or session['chef_naam'] != chef_naam:
        flash("Geen toegang. Log opnieuw in.", "danger")
        return redirect(url_for('login'))
    
    conn = get_db_connection()
    if conn is None:
        flash("Database connection error.", "danger")
        return redirect(url_for('takenboek.index', chef_naam=chef_naam))
    
    cur = conn.cursor(dictionary=True)
    
    try:
        # Get all tasks for this chef - Fix the table name here
        cur.execute("""
            SELECT * FROM student_tasks
            WHERE chef_id = %s
            ORDER BY blok, onderdeel, techniek
        """, (session['chef_id'],))
        tasks = cur.fetchall()
        
        if not tasks:
            flash("Er zijn geen taken om te exporteren.", "warning")
            return redirect(url_for('takenboek.index', chef_naam=chef_naam))
        
        # Create Word document
        doc = Document()
        doc.add_heading(f'Takenboek - {chef_naam}', 0)
        
        # Add table with tasks
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Blok"
        header_cells[1].text = "Onderdeel"
        header_cells[2].text = "Techniek"
        header_cells[3].text = "Taak"
        header_cells[4].text = "Uitgevoerd op"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add tasks to table
        for task in tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = task['blok']
            row_cells[1].text = task['onderdeel']
            row_cells[2].text = task['techniek']
            row_cells[3].text = task['taak']
            row_cells[4].text = task['uitgevoerd_op'].strftime('%d-%m-%Y') if task['uitgevoerd_op'] else '-'
            
            # Add comments if available
            if task['opmerkingen']:
                if task['opmerkingen'].strip():
                    comment_row = table.add_row().cells
                    comment_row[0].merge(comment_row[4])
                    comment_row[0].text = f"Opmerkingen: {task['opmerkingen']}"
        
        # Create in-memory file
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Send file to user
        return send_file(
            buffer,
            as_attachment=True,
            download_name='takenboek.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f'Error exporting tasks: {str(e)}')
        flash(f"Er is een fout opgetreden bij het exporteren: {str(e)}", "danger")
        return redirect(url_for('takenboek.index', chef_naam=chef_naam))
    
    finally:
        cur.close()
        conn.close()
